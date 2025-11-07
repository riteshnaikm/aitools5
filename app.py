import os
import pdfplumber
import logging
import hashlib
import json
import nltk
import sqlite3
from nltk.tokenize import sent_tokenize
from flask import Flask, request, jsonify, render_template, Response, stream_with_context, make_response
from dotenv import load_dotenv
from pinecone import Pinecone, ServerlessSpec
from langchain_community.vectorstores import Pinecone as PineconeVectorStore
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from rank_bm25 import BM25Okapi
from functools import lru_cache
import re
import pandas as pd
import warnings
import google.generativeai as genai
from docx import Document
import uuid
from werkzeug.utils import secure_filename
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor
from asgiref.wsgi import WsgiToAsgi
import time
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
import markdown
from html import unescape


# Suppress specific warnings
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! top_p is not default parameter.")
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! presence_penalty is not default parameter.")
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! frequency_penalty is not default parameter.")

# Load environment variables
load_dotenv()

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_INDEX_NAME = "hr-knowledge-base"
POLICIES_FOLDER = "HR_docs/"
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Add this dictionary after imports
ACRONYM_MAP = {
    "wfh": "work from home policy",
    "pto": "paid time off policy",
    "loa": "leave of absence policy",
    "nda": "non-disclosure agreement",
    "od": "on duty policy",
    "hrbp": "human resources business partner",
    "kra": "KRA Policy - Promoting Transparency",
    "regularization": "Time change Request/ Regularization",
    "regularisation": "Time change Request/ Regularization",
    "posh": "Policy On Prevention of Sexual Harassment",
    "appraisal": "PERFORMANCE APPRAISAL & PROMOTION POLICY",
    "promotion": "PERFORMANCE APPRAISAL & PROMOTION POLICY",
    "prep": "Performance Review & Enhancement Program",
    "Grade": "GRADE STRUCTURE & FLEXIBILITY",
    "leave": "LEAVE POLICY",
    "nda": "Non Compete and Non Disclosure",
    "Office timings": "Office Timing and Attendance Policy",
    "pet": "pet policy",
    "sprint": "Weekly Sprint Policy",
    "work ethics": "WORK PLACE ETHICS"
}

# Standard behavioral questions
QUICK_CHECKS = [
    "Are you willing to relocate if applicable?",
    "What is your notice period?",
    "Can you provide details about your current organization?",
    "Please describe your current role and responsibilities.",
    "What is your current CTC (Cost to Company)?",
    "What is your expected CTC?",
    "What is your educational background?",
    "Can you describe any significant projects you've worked on?",
    "Are there any specific client requirements you want to discuss?",
    "Do you have references from colleagues who might be interested in opportunities with us?"
]

# Prompt templates
input_prompt_template = """
Act as a highly skilled ATS (Applicant Tracking System) specializing in evaluating resumes for job descriptions provided. Your information will be consumed by fellow HR professionals to help them evaluate resumes quickly.

### Task:
Evaluate the provided **resume** against the given **job description**. Consider industry trends and the competitive job market. Prioritize factors based on their relevance to the specific requirements of the job description. All Match Factors should be weighted equally unless otherwise specified in the job description.

### Certification Handling:
* If the job description explicitly mentions required certifications, score the candidate based on whether they possess those certifications.
* If the job description does not mention certifications, consider any relevant certifications the candidate has as a potential bonus, but do not penalize candidates who lack them.
* If a candidate lacks a certification that is explicitly mentioned in the job description, lower the overall score significantly.

### Overqualification Analysis:
* Carefully assess if the candidate is **overqualified or "overkill"** for the position.
* Even if all criteria match perfectly, evaluate if the candidate's experience level, seniority, or qualifications significantly exceed the role requirements.
* Consider: years of experience (e.g., 15 years for a 3-5 year role), job title level (e.g., VP/Director applying for mid-level), salary expectations, and potential flight risk.
* Flag candidates who may be too senior and likely to leave quickly or become disengaged.

### Output:
Return a valid JSON object ONLY. The JSON object MUST have the following keys:

* `"JD Match"` (string): Percentage match (e.g., "85%").
* `"MissingKeywords"` (list): List of missing keywords (can be empty).
* `"Profile Summary"` (string): Summary of strengths and areas for improvement.
* `"Over/UnderQualification Analysis"` (string): **CRITICAL** - Assess if candidate is overqualified, underqualified, or right-fit. Include: (1) Is candidate too senior OR too junior for this role? (2) Experience mismatch (years, title level)? (3) Likely salary expectations vs budget? (4) Flight risk or capability gap assessment? (5) Recommendation: "Perfect fit" OR "Overqualified - High flight risk" OR "Underqualified - needs development" OR "Overqualified but worth considering" with clear reasoning.
* `"Match Factors"` (object): Breakdown of factors that contributed to the match percentage with individual scores:
    * `"Skills Match"` (number): 0-100 score for technical skills alignment
    * `"Experience Match"` (number): 0-100 score for experience level alignment
    * `"Education Match"` (number): 0-100 score for education requirements match
    * `"Industry Knowledge"` (number): 0-100 score for relevant industry knowledge
    * `"Certification Match"` (number): 0-100 score for relevant certifications
* `"Reasoning"` (string): Explanation of the scoring decision for each "Match Factor" and the overall "JD Match" score.
* `"Candidate Fit Analysis"` (object): **NEW** - Detailed recruiter-ready fit analysis with the following structure:
    * `"Dimension Evaluation"` (array of objects): Each object contains:
        - `"Dimension"` (string): e.g., "Overall Profile Match", "Industry Exposure", "Hands-on Technical Depth", etc. - Only include dimensions RELEVANT to this specific JD.
        - `"Evaluation"` (string): e.g., "✅ Strong", "✅ Good", "✅ Excellent", "✅ Outstanding", "⚠️ Potential risk area", "⚠️ Moderate", "❌ Weak"
        - `"Recruiter Comments"` (string): Detailed, specific evidence from resume explaining the evaluation
    * `"Risk and Gaps"` (array of objects or null): Each object contains:
        - `"Area"` (string): The risk area
        - `"Risk"` (string): Description of the risk
        - `"Recruiter Strategy"` (string): Specific question or approach to probe this in screening
      If NO major risks, set this to null or empty array.
    * `"Recommendation"` (object):
        - `"Verdict"` (string): e.g., "✅ Strong Shortlist", "⚠️ Conditional Shortlist", "❌ Not Recommended"
        - `"Fit Level"` (string): e.g., "9/10", "High", "Medium-High", etc.
        - `"Rationale"` (string): 2-4 sentence summary of why this verdict was reached
    * `"Recruiter Narrative"` (string): A ready-to-submit 2-5 sentence paragraph that recruiters can use when presenting this candidate to hiring managers. Should highlight: key strengths, years of experience, industry relevance, location (if mentioned), and why they're a good fit.

Do NOT include any additional text, explanations, or formatting outside the JSON object.

---
**Resume:** {resume_text}
**Job Description:** {job_description}
"""

interview_questions_prompt = """
You are an experienced technical recruiter preparing for an interview. Based on the candidate's resume and the job description, generate relevant interview questions.

### Task:
Generate two sets of interview questions - 10 technical questions and 10 non-technical questions that are specifically tailored to assess this candidate for this role.

### Output:
Return a valid JSON object ONLY with the following keys:
* `"TechnicalQuestions"` (array): 10 technical questions related to the candidate's skills and the job requirements.
* `"NonTechnicalQuestions"` (array): 10 behavioral, situational, or cultural fit questions.

Each question should be thoughtful, specific to the resume and job description, and reveal important information about the candidate's suitability.

Do NOT include any additional text, explanations, or formatting outside the JSON object.

---
**Resume:** {resume_text}
**Job Description:** {job_description}
**Candidate Profile Summary:** {profile_summary}
"""

job_stability_prompt = """
As an HR analytics expert, analyze the work history in this resume to determine if the candidate shows job-hopping tendencies.

### Task:
Review the resume and identify the candidate's job history, analyzing tenure at each position to evaluate stability.

### Output:
Return a valid JSON object ONLY with the following keys:
* `"IsStable"` (boolean): true if candidate shows good job stability, false if there are job-hopping concerns
* `"AverageJobTenure"` (string): estimated average time spent at each position (e.g., "2.5 years")
* `"JobCount"` (number): total number of positions held
* `"StabilityScore"` (number): 0-100 score indicating job stability (higher is better)
* `"ReasoningExplanation"` (string): brief explanation of the stability assessment
* `"RiskLevel"` (string): "Low", "Medium", or "High" risk of leaving quickly

Do NOT include any additional text, explanations, or formatting outside the JSON object.

---
**Resume:** {resume_text}
"""


# Add career progression prompt template after other prompt templates
career_prompt = """
You are an expert HR analyst. Analyze the candidate's career progression from their resume.
Focus on identifying career growth, job transitions, and progression patterns.

Provide your analysis in the following JSON format ONLY:
{
    "red_flags": [
        "<potential concern about career progression>",
        "<another concern if any>"
    ],
    "reasoning": [
        "<bullet point about career progression>",
        "<another key insight about their career trajectory>",
        "<pattern observed in their job transitions>"
    ]
}

Guidelines:
- red_flags should identify potential concerns for hiring (job hopping, gaps, demotions, etc.). Can be empty array if no major concerns.
- reasoning should be an ARRAY of bullet points (not a paragraph) explaining the overall career progression analysis, highlighting growth patterns, stability, and trajectory.

Resume text:
{resume_text}
"""

# Recruiter Handbook Prompt Template
recruiter_handbook_prompt = """
SYSTEM:
You are an expert technical recruiter and talent evaluator specializing in AI, analytics, product, and consulting roles. 
Your task is to generate a detailed recruiter-style evaluation and fit analysis based on a Job Description (JD) and a Candidate Resume. 
You must think and write like a senior talent partner at a top-tier consulting firm (Fractal, Deloitte, BCG, etc.) — structured, insightful, data-driven, and nuanced. 
Your output will form a Recruiter Handbook that helps internal recruiters, interviewers, and hiring managers make data-driven shortlisting decisions.

---

INSTRUCTIONS:

You will be provided with:
- JOB_DESCRIPTION_TEXT: {job_description}
- CANDIDATE_RESUME_TEXT: {resume_text}

Generate the following sections in professional markdown format:

### 1️⃣ JD Snapshot
Summarize the role in 5–6 lines:
- Role title, level, domain/vertical
- Key skills required
- Nature of role (hands-on, consulting-led, AI-driven, etc.)
- Success indicators

---

### 2️⃣ Candidate Summary
Summarize the candidate's profile in 5–6 lines:
- Education and total experience
- Functional & technical focus areas
- Domain experience (industries)
- Career trajectory highlights
- Distinguishing achievements

---

### 3️⃣ Fit Matrix (Comparison Table)
Create a markdown table with columns:  
**Dimension | JD Expectation | Candidate Evidence | Rating (1–5) | Comment**

Include these dimensions:
- Domain Expertise  
- Technical / AI / Cloud Engineering Depth  
- Consulting Gravitas / CXO Advisory Experience  
- Innovation & IP / Emerging Tech Thought Leadership  
- Hands-on IC Credibility  
- Project & Delivery Leadership  
- Business Acumen & Commercial Awareness  
- Culture / Communication Fit  

Rating legend: 5=Excellent | 3=Average | 1=Weak

---

### 4️⃣ Detailed Fit Commentary
Write 5–8 paragraphs of nuanced recruiter commentary:
- Where the candidate aligns strongly  
- Where they are weak or untested  
- How their consulting/technical balance fits  
- Domain alignment and value they could bring  
- Potential role re-alignment if not an exact fit  
Tone: confident, analytical, and evidence-driven.

---

### 5️⃣ Red Flags & Mitigation
List 3–5 potential red flags (if any) and ways to mitigate or probe during interview.

---

### 6️⃣ Interview Focus Areas
List 8–10 recommended interview questions, grouped under:
- Technical/Engineering depth  
- Consulting & stakeholder management  
- Domain expertise  
- AI/GenAI awareness  
- Leadership & delivery

---

### 7️⃣ Recruiter Summary & Pitch
Write 2 short paragraphs (recruiter voice) that summarize:
- Why this candidate could be compelling to the client
- How to position them internally
Use persuasive, professional tone for internal submission.

---

### 8️⃣ Final Verdict
Provide:
- **Fit Verdict:** Strong Shortlist / Conditional / Reconsider / Reject  
- **Fit Score (0–100%):**  
- **Confidence (High / Medium / Low):**  
- **Best-fit Domain (Retail / FSI / TMT / HLS / Other):**

---

**Output Formatting:**
- Write the recruiter handbook in professional markdown (for rendering in a web UI).  
- Keep tone polished, confident, and analytical — like a top-tier recruiter brief.  
- Be specific, never generic. Use evidence-based phrases like:
  - "Shows strong multi-client consulting maturity."
  - "Demonstrates architectural thought leadership but lacks GenAI delivery exposure."
  - "Consulting gravitas evident from CIO-level advisory roles."
- Do NOT include markdown code block markers in your response.
"""

# Initialize Gemini model
genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
model = genai.GenerativeModel('gemini-2.0-flash')

# Initialize Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
asgi_app = WsgiToAsgi(app)

# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize NLTK
nltk.download("punkt")


# Initialize Groq LLM
llm = ChatGroq(
    groq_api_key=GROQ_API_KEY,
    # model_name="mixtral-8x7b-32768",  # This generates long text,  max_tokens=4096
    # model_name=   "llama-3.1-8b-instant",#"qwen-2.5-32b", #"deepseek-r1-distill-qwen-32b",
    model_name = "qwen/qwen3-32b",
    temperature=0.377,
    max_tokens=2048,   #4096
    top_p=0.95,
    presence_penalty=0.1,
    frequency_penalty=0.1
)

# Initialize Pinecone
pc = Pinecone(api_key=PINECONE_API_KEY)
index_name = PINECONE_INDEX_NAME
if index_name not in pc.list_indexes().names():
    pc.create_index(
        name=index_name,
        dimension=384,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )
index = pc.Index(index_name)

# Initialize embeddings
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Initialize vector store
vectorstore = None
try:
    from langchain_pinecone import PineconeVectorStore as NewPineconeVectorStore
    vectorstore = NewPineconeVectorStore(
        index=index,
        embedding=embeddings,
        text_key="text"
    )
    logging.info("✅ Using new langchain-pinecone vectorstore")
except ImportError:
    # Fallback to old import if new package not available
    try:
        from langchain_community.vectorstores import Pinecone as PineconeVectorStore
        vectorstore = PineconeVectorStore(
            index=index,
            embedding=embeddings,
            text_key="text"
        )
        logging.info("✅ Using old langchain-community vectorstore")
    except Exception as e:
        logging.error(f"❌ Error initializing vectorstore: {e}")
        vectorstore = None
except Exception as e:
    logging.error(f"❌ Error initializing new vectorstore: {e}")
    vectorstore = None

# Initialize database
DATABASE_NAME = 'combined_db.db'

def init_db():
    """Initialize database with all required tables"""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    
    # Create evaluations table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS evaluations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            resume_path TEXT,
            filename TEXT,
            job_title TEXT,
            job_description TEXT,
            match_percentage REAL,
            match_factors TEXT,
            profile_summary TEXT,
            missing_keywords TEXT,
            job_stability TEXT,
            career_progression TEXT,
            technical_questions TEXT,
            nontechnical_questions TEXT,
            behavioral_questions TEXT,
            oorwin_job_id TEXT,
            candidate_fit_analysis TEXT,
            over_under_qualification TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create qa_history table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS qa_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question TEXT NOT NULL,
            retrieved_docs TEXT,
            final_answer TEXT,
            feedback TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create qa_feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS qa_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER,
            rating INTEGER,
            feedback TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(question_id),
            FOREIGN KEY (question_id) REFERENCES qa_history (id)
        )
    ''')
    
    # Create feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            evaluation_id INTEGER,
            rating INTEGER,
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(evaluation_id),
            FOREIGN KEY (evaluation_id) REFERENCES evaluations (id)
        )
    ''')
    
    # Create handbook_feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS handbook_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            handbook_id INTEGER NOT NULL,
            rating INTEGER NOT NULL CHECK(rating >= 1 AND rating <= 5),
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(handbook_id),
            FOREIGN KEY (handbook_id) REFERENCES recruiter_handbooks (id)
        )
    ''')
    
    # Create interview_questions table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS interview_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            evaluation_id INTEGER,
            technical_questions TEXT,
            nontechnical_questions TEXT,
            behavioral_questions TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (evaluation_id) REFERENCES evaluations (id)
        )
    ''')
    
    # Create recruiter_handbooks table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS recruiter_handbooks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            oorwin_job_id TEXT,
            job_title TEXT,
            job_description TEXT,
            additional_context TEXT,
            markdown_content TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Handle schema updates for existing tables
    try:
        # Check if evaluations table has new columns
        cursor.execute("PRAGMA table_info(evaluations)")
        columns = [col[1] for col in cursor.fetchall()]
        
        if 'candidate_fit_analysis' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN candidate_fit_analysis TEXT')
            print("Added candidate_fit_analysis column to evaluations")
        
        if 'over_under_qualification' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN over_under_qualification TEXT')
            print("Added over_under_qualification column to evaluations")
    except Exception as e:
        print(f"Note: Schema update check: {e}")
    
    conn.commit()
    conn.close()

# Initialize database at startup
init_db()

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_gemini_response(input_prompt):
    """Get response from Gemini model and clean it up."""
    try:
        response = model.generate_content(input_prompt)
        response_text = response.text.strip()
        
        # Remove markdown code block markers if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
            
        # Clean up any extra whitespace and newlines
        response_text = response_text.strip()
        
        # Try to parse as JSON to validate
        try:
            parsed_json = json.loads(response_text)
            return json.dumps(parsed_json)  # Return properly formatted JSON string
        except json.JSONDecodeError:
            # If not valid JSON, try to extract JSON using regex
            match = re.search(r"\{.*\}", response_text, re.DOTALL)
            if match:
                try:
                    parsed_json = json.loads(match.group(0))
                    return json.dumps(parsed_json)  # Return properly formatted JSON string
                except json.JSONDecodeError:
                    raise ValueError("Invalid JSON structure in response")
            else:
                raise ValueError("No valid JSON found in response")
                
    except Exception as e:
        logging.error(f"Error in get_gemini_response: {str(e)}")
        return json.dumps({})  # Return valid empty JSON object as fallback

def extract_text_from_file(file_path):
    try:
        ext = file_path.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ModuleNotFoundError as e:
                if "PyCryptodome" in str(e) or "Crypto" in str(e):
                    return None, "PyCryptodome is required for encrypted PDFs. Please install it with 'pip install pycryptodome'."
                raise
        elif ext == 'docx':
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        elif ext == 'doc':
            return None, "Support for .doc files is limited. Please convert to .docx or PDF."
        else:
            return None, "Unsupported file format."
    except Exception as e:
        logging.error(f"File extraction error: {str(e)}")
        return None, str(e)

def hybrid_search(query, k=5):
    """Perform hybrid search using BM25 and vector similarity."""
    try:
        # Get vector search results
        vector_results = vectorstore.similarity_search(query, k=k)
        
        # Get BM25 results
        bm25_results = []
        if os.path.exists(POLICIES_FOLDER):
            for filename in os.listdir(POLICIES_FOLDER):
                if filename.endswith(('.txt', '.md')):
                    with open(os.path.join(POLICIES_FOLDER, filename), 'r', encoding='utf-8') as f:
                        text = f.read()
                        sentences = sent_tokenize(text)  # Use NLTK sentence tokenizer
                        bm25_results.extend(sentences)
        
        # Combine and deduplicate results
        combined_results = []
        seen_texts = set()
        
        # Add vector search results
        for doc in vector_results:
            if doc.page_content not in seen_texts:
                combined_results.append(doc.page_content)
                seen_texts.add(doc.page_content)
        
        # Add BM25 results
        for sentence in bm25_results:
            if sentence not in seen_texts:
                combined_results.append(sentence)
                seen_texts.add(sentence)
        
        # Join results with newlines
        return "\n".join(combined_results)
    
    except Exception as e:
        logging.error(f"Error in hybrid_search: {e}")
        return ""

def save_evaluation(eval_id, filename, job_title, rank_score, missing_keywords, profile_summary, match_factors, job_stability, additional_info=None, oorwin_job_id=None, candidate_fit_analysis=None, over_under_qualification=None):
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Convert data to JSON strings if they're not already
        try:
            missing_keywords_json = json.dumps(missing_keywords) if not isinstance(missing_keywords, str) else missing_keywords
        except Exception as e:
            logging.error(f"Error converting missing_keywords to JSON: {e}")
            missing_keywords_json = '[]'
            
        try:
            match_factors_json = json.dumps(match_factors) if not isinstance(match_factors, str) else match_factors
        except Exception as e:
            logging.error(f"Error converting match_factors to JSON: {e}")
            match_factors_json = '{}'
            
        try:
            job_stability_json = json.dumps(job_stability) if not isinstance(job_stability, str) else job_stability
        except Exception as e:
            logging.error(f"Error converting job_stability to JSON: {e}")
            job_stability_json = '{}'
        
        # Ensure all JSON strings are valid
        if not missing_keywords_json or missing_keywords_json == 'null':
            missing_keywords_json = '[]'
        if not match_factors_json or match_factors_json == 'null':
            match_factors_json = '{}'
        if not job_stability_json or job_stability_json == 'null':
            job_stability_json = '{}'
        
        # Convert rank_score to integer if it's a string
        rank_score_int = int(rank_score) if isinstance(rank_score, str) else rank_score
        
        # Ensure all values are strings except rank_score_int and eval_id
        filename_str = str(filename)
        job_title_str = str(job_title)
        profile_summary_str = str(profile_summary)
        
        # Handle oorwin_job_id (can be None or empty string)
        oorwin_job_id_str = str(oorwin_job_id).strip() if oorwin_job_id else None
        if oorwin_job_id_str == '' or oorwin_job_id_str == 'None':
            oorwin_job_id_str = None
        
        # Convert additional_info to JSON string if it's a dict or list
        if isinstance(additional_info, (dict, list)):
            additional_info_str = json.dumps(additional_info)
        else:
            additional_info_str = str(additional_info) if additional_info is not None else ""
        
        # Extract career progression from additional_info
        career_progression = additional_info.get('career_progression', {}) if isinstance(additional_info, dict) else {}
        career_progression_json = json.dumps(career_progression)
        
        # Convert new fields to JSON
        candidate_fit_analysis_json = json.dumps(candidate_fit_analysis) if candidate_fit_analysis else '{}'
        over_under_qualification_str = str(over_under_qualification) if over_under_qualification else ''
        
        # Debug: Log the actual values being inserted
        logging.info(f"Values to insert - eval_id: {eval_id}, filename: {filename_str}, job_title: {job_title_str}")
        logging.info(f"JSON values - missing_keywords_json type: {type(missing_keywords_json)}, value: {missing_keywords_json[:100] if len(missing_keywords_json) > 100 else missing_keywords_json}")
        logging.info(f"JSON values - match_factors_json type: {type(match_factors_json)}")
        logging.info(f"All param types: rank_score_int={type(rank_score_int)}, oorwin_job_id_str={type(oorwin_job_id_str)}, datetime={type(datetime.now())}")
        
        # Convert datetime to string
        timestamp_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Prepare all parameters as a tuple
        params = (
            eval_id, filename_str, filename_str, job_title_str, "", rank_score_int, 
            match_factors_json, profile_summary_str, missing_keywords_json, 
            job_stability_json, career_progression_json, None, None, None, oorwin_job_id_str, 
            candidate_fit_analysis_json, over_under_qualification_str, timestamp_str
        )
        
        # Log all parameter types
        logging.info(f"Parameter types: {[type(p).__name__ for p in params]}")
        
        cursor.execute(
            """
            INSERT INTO evaluations (
                resume_path, filename, job_title, job_description, match_percentage, 
                match_factors, profile_summary, missing_keywords, 
                job_stability, career_progression, technical_questions,
                nontechnical_questions, behavioral_questions, oorwin_job_id, 
                candidate_fit_analysis, over_under_qualification, timestamp
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            params[1:]  # Skip the first parameter (eval_id)
            )
        
        # Get the auto-generated ID
        db_id = cursor.lastrowid
        
        conn.commit()
        
        # Log successful save
        logging.info(f"✅ EVALUATION SAVED TO DATABASE!")
        logging.info(f"   database_id: {db_id}")
        logging.info(f"   filename: {filename_str}")
        logging.info(f"   job_title: {job_title_str}")
        logging.info(f"   match_percentage: {rank_score_int}")
        logging.info(f"   oorwin_job_id: {oorwin_job_id_str}")
        
        conn.close()
        return db_id  # Return the database ID instead of True
    except Exception as e:
        logging.error(f"Database error in save_evaluation: {str(e)}", exc_info=True)
        logging.error(f"Data being saved - eval_id: {eval_id}, filename: {filename}, job_title: {job_title}")
        logging.error(f"Data types - rank_score: {type(rank_score)}, missing_keywords: {type(missing_keywords)}")
        return False

def save_feedback(evaluation_id, rating, comments):
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO feedback (evaluation_id, rating, comments, timestamp) VALUES (?, ?, ?, ?)",
            (evaluation_id, rating, comments, datetime.now())
        )
        logging.debug(f"Feedback inserted: evaluation_id={evaluation_id}, rating={rating}, comments={comments}")
        conn.commit()
        conn.close()
        return True
    except sqlite3.Error as e:
        logging.error(f"Database error in save_feedback: {str(e)}")
        return False
    except Exception as e:
        logging.error(f"Unexpected error in save_feedback: {str(e)}")
        return False

def save_interview_questions(evaluation_id, technical_questions, nontechnical_questions, behavioral_questions):
    """Save interview questions to database"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        technical_json = json.dumps(technical_questions) if not isinstance(technical_questions, str) else technical_questions
        nontechnical_json = json.dumps(nontechnical_questions) if not isinstance(nontechnical_questions, str) else nontechnical_questions
        behavioral_json = json.dumps(behavioral_questions) if not isinstance(behavioral_questions, str) else behavioral_questions
        
        cursor.execute(
            "INSERT INTO interview_questions (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions, timestamp) VALUES (?, ?, ?, ?, ?)",
            (evaluation_id, technical_json, nontechnical_json, behavioral_json, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        )
        conn.commit()
        conn.close()
        logging.debug(f"Interview questions saved successfully for: {evaluation_id}")
        return True
    except Exception as e:
        logging.error(f"Database error in save_interview_questions: {str(e)}")
        return False

def save_recruiter_handbook(evaluation_id, markdown_content, json_summary):
    """Save recruiter handbook to database"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Ensure json_summary is a string
        json_summary_str = json_summary if isinstance(json_summary, str) else json.dumps(json_summary)
        
        cursor.execute(
            "INSERT INTO recruiter_handbooks (evaluation_id, markdown_content, json_summary, timestamp) VALUES (?, ?, ?, ?)",
            (evaluation_id, markdown_content, json_summary_str, datetime.now())
        )
        conn.commit()
        conn.close()
        logging.debug(f"Recruiter handbook saved successfully for: {evaluation_id}")
        return True
    except Exception as e:
        logging.error(f"Database error in save_recruiter_handbook: {str(e)}")
        return False

# Add these constants near the top with other constants
BOT_INFO = {
    "name": "PeopleBot",
    "creator": "PeopleLogic",
    "responsibility": "Help recruiters in HR policies, benefits & with any other questions!",
    "capabilities": "Help recruiters in HR policies, benefits & with any other questions"
}

GREETINGS = ["hello", "hi", "hey", "greetings", "good morning", "good afternoon", "good evening"]
IDENTITY_QUESTIONS = [
    "who are you",
    "who are u",
    "who r yu",
    "who r u",
    "who are you",
    "what is your name",
    "what are you",
    "who built you",
    "what are u",
    "who built u",
    "who created you",
    "what can you do",
    "what do you do",
    "what is your name",
    "tell me about yourself",
    "tell me about you",
     "Who Are You", "WHO ARE YOU", "Who r U", "wHo bUiLt YoU",  
    "WHAT IS YOUR NAME", "wHaT dO yOu dO",  "whoz u", "what u do", "wots ur name", "whut do u do",  
    "wats ur function", "who dun made u",  "what are you capable of", "what skills do you have",  
    "describe yourself", "what is your role", "explain yourself",  
    "what services do you provide", "how do you work",  
    "can you tell me what you do", "can you tell me your name", "can you introduce yourself",  
    "what should I call you", "how do I address you",  
    "could you tell me what you are", "what do people call you",  
    "what is your full name", "how were you created",  
    "give me some details about you", "tell me your background",  "who u", "who dis", "whats ur deal", "whats ur function", "who tf r u",  
    "who's this", "u bot?", "who you be", "ur name?", "who made u",  
    "who's your maker", "who made this bot", "whu r u", "whi r u", "who r yu", "wht is ur name", "whts ur name",  
    "whats ur name", "wat is your name", "wat r u", "whu built u",  
    "who made you", "who designed you", "who programmed you", "what are you capable of", "what skills do you have",  
    "describe yourself", "what is your role", "explain yourself",  
    "what services do you provide", "how do you work",  
    "can you tell me what you do",  

    # Mixed uppercase/lowercase variations  
    "Who Are You", "WHO ARE YOU", "Who r U", "wHo bUiLt YoU",  
    "WHAT IS YOUR NAME", "wHaT dO yOu dO",  

    # Slang & misspellin
    "whoz u", "what u do", "wots ur name", "whut do u do",  
    "wats ur function", "who dun made u",  

    # Phonetic spellings & accents  
    "hoo r u", "wat iz ur name", "hoo maid u", "wat cn u do",  
    "whut ur name", "whachu do",  

    # Extra punctuation variations  
    "who are you?", "who are you!", "who are you??",  
    "what is your name?", "who built you??", "who made you?!",
]

def handle_special_queries(question):
    """Handle greetings and identity-related questions."""
    question_lower = question.lower().strip("?!. ")
    
    # Handle greetings
    if question_lower in GREETINGS:
        return f"Hello! I'm {BOT_INFO['name']}, your HR assistant. How can I help you today?"
    
    # Handle identity questions
    if any(q in question_lower for q in IDENTITY_QUESTIONS):
        if "who" in question_lower or "what is your name" in question_lower:
            return f"I'm {BOT_INFO['name']}, an AI assistant built by {BOT_INFO['creator']}. {BOT_INFO['responsibility']}"
        elif "created" in question_lower or "built" in question_lower:
            return f"I was created by {BOT_INFO['creator']} to {BOT_INFO['responsibility']}"
        elif "can you do" in question_lower or "do you do" in question_lower:
            return f"I can {BOT_INFO['capabilities']}"
        else:
            return f"I'm {BOT_INFO['name']}, an AI assistant created by {BOT_INFO['creator']}. {BOT_INFO['capabilities']}"

    # Handle holiday list queries (static 2025 list provided by HR)
    if "holiday" in question_lower or "holidays" in question_lower:
        year = "2025"
        header = f"## Company Holidays {year}\n\nBelow are the declared holidays for {year}.\n\n"

        india_table = (
            "### India Offices (Bangalore/APAC & EU, Hyderabad, Mumbai, Delhi)\n"
            "| Date | Day | Bangalore/APAC & EU | Hyderabad | Mumbai | Delhi |\n"
            "|------|-----|----------------------|-----------|--------|-------|\n"
            "| 1-Jan-2025 | Wednesday | New Year | New Year | New Year | New Year |\n"
            "| 14-Jan-2025 | Tuesday | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti |\n"
            "| 14-Mar-2025 | Friday | - | Holi | Holi | Holi |\n"
            "| 31-Mar-2025 | Monday | Ramzan (Id Ul Fitr) | Ramzan (Id Ul Fitr) | Ramzan (Id Ul Fitr) | - |\n"
            "| 18-Apr-2025 | Friday | Good Friday | - | - | Good Friday |\n"
            "| 1-May-2025 | Thursday | May Day | May Day | May Day | May Day |\n"
            "| 15-Aug-2025 | Friday | Independence Day | Independence Day | Independence Day | Independence Day |\n"
            "| 27-Aug-2025 | Wednesday | Ganesh Chaturthi | Ganesh Chaturthi | Ganesh Chaturthi | Ganesh Chaturthi |\n"
            "| 2-Oct-2025 | Thursday | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara |\n"
            "| 20-Oct-2025 | Monday | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi |\n"
            "| 25-Dec-2025 | Thursday | Christmas | Christmas | Christmas | Christmas |\n"
        )

        us_table = (
            "\n### Global Services - US\n"
            "| Date | Day | Holiday |\n"
            "|------|-----|---------|\n"
            "| 1-Jan-2025 | Wednesday | New Year |\n"
            "| 18-Apr-2025 | Friday | Good Friday |\n"
            "| 26-May-2025 | Monday | Memorial Day |\n"
            "| 4-Jul-2025 | Friday | Independence Day |\n"
            "| 1-Sep-2025 | Monday | Labour Day |\n"
            "| 20-Oct-2025 | Monday | Diwali |\n"
            "| 27-Nov-2025 | Thursday | Thanksgiving |\n"
            "| 28-Nov-2025 | Friday | Day after Thanksgiving |\n"
            "| 24-Dec-2025 | Wednesday | Christmas Eve |\n"
            "| 25-Dec-2025 | Thursday | Christmas Day |\n"
        )

        footnote = "\n> Note: If a holiday falls on a weekend, local HR guidelines on compensatory off apply."
        return header + india_table + us_table + footnote
    
    return None

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/hr-assistant')
def hr_assistant():
    return render_template('index1.html')

@app.route('/resume-evaluator')
def resume_evaluator():
    return render_template('index2.html')

@app.route('/test-tabs')
def test_tabs():
    return render_template('test_tabs.html')

@app.route('/evaluation/<int:eval_id>')
def view_evaluation(eval_id):
    """View a single evaluation in detail"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Fetch evaluation data
        cursor.execute('''
            SELECT 
                e.id, e.filename, e.job_title, e.job_description,
                e.match_percentage, e.match_factors, e.profile_summary,
                e.missing_keywords, e.job_stability, e.career_progression,
                e.oorwin_job_id, e.timestamp
            FROM evaluations e
            WHERE e.id = ?
        ''', (eval_id,))
        
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return "Evaluation not found", 404
        
        # Fetch interview questions
        cursor.execute('''
            SELECT technical_questions, nontechnical_questions, behavioral_questions
            FROM interview_questions
            WHERE evaluation_id = ?
        ''', (eval_id,))
        
        questions_row = cursor.fetchone()
        conn.close()
        
        # Parse JSON fields
        import json
        evaluation = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'job_description': row[3],
            'match_percentage': row[4],
            'match_factors': json.loads(row[5]) if row[5] else {},
            'profile_summary': row[6],
            'missing_keywords': json.loads(row[7]) if row[7] else [],
            'job_stability': json.loads(row[8]) if row[8] else {},
            'career_progression': json.loads(row[9]) if row[9] else {},
            'oorwin_job_id': row[10],
            'timestamp': row[11]
        }
        
        if questions_row:
            evaluation['technical_questions'] = json.loads(questions_row[0]) if questions_row[0] else []
            evaluation['nontechnical_questions'] = json.loads(questions_row[1]) if questions_row[1] else []
            evaluation['behavioral_questions'] = json.loads(questions_row[2]) if questions_row[2] else []
        else:
            evaluation['technical_questions'] = []
            evaluation['nontechnical_questions'] = []
            evaluation['behavioral_questions'] = []
        
        return render_template('evaluation_view.html', evaluation=evaluation)
        
    except Exception as e:
        logging.error(f"Error viewing evaluation {eval_id}: {str(e)}")
        return f"Error loading evaluation: {str(e)}", 500

@app.route('/history')
def history():
    conn = sqlite3.connect('combined_db.db')
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
            SELECT 
                e.id, 
                e.filename, 
                e.job_title, 
                e.match_percentage, 
                e.missing_keywords, 
                e.profile_summary, 
                e.job_stability,
                e.career_progression,
                e.timestamp,
                iq.technical_questions,
                iq.nontechnical_questions,
                e.oorwin_job_id
            FROM evaluations e
            LEFT JOIN interview_questions iq ON e.id = iq.evaluation_id
            ORDER BY e.timestamp DESC
        ''')
        
        evaluations = []
        for row in cursor.fetchall():
            try:
                # Helper function for safe JSON parsing
                def safe_json_loads(data, default):
                    if not data:
                        logging.info(f"Empty data for field, using default: {default}")
                        return default
                    try:
                        if isinstance(data, str):
                            return json.loads(data)
                        return data
                    except json.JSONDecodeError as e:
                        logging.error(f"JSON parsing error for evaluation {row[0]}: {str(e)} - Data: {data}")
                        # Try to clean the string if it's malformed
                        if isinstance(data, str):
                            try:
                                # Remove any trailing commas, fix quotes
                                cleaned = re.sub(r',\s*}', '}', data)
                                cleaned = re.sub(r',\s*]', ']', cleaned)
                                return json.loads(cleaned)
                            except:
                                pass
                        return default
                
                # Parse JSON fields with robust error handling
                missing_keywords_raw = row[4]
                try:
                    if missing_keywords_raw:
                        missing_keywords = safe_json_loads(missing_keywords_raw, [])
                        # If it's a string that looks like a list but isn't parsed as one
                        if not isinstance(missing_keywords, list):
                            # Try to extract keywords from a string representation
                            if isinstance(missing_keywords, str):
                                # Remove brackets and split by commas
                                missing_keywords = [k.strip(' "\'') for k in missing_keywords.strip('[]').split(',')]
                            else:
                                missing_keywords = [str(missing_keywords)]
                    else:
                        missing_keywords = []
                except Exception as e:
                    logging.error(f"Error parsing missing_keywords for eval {row[0]}: {str(e)}")
                    missing_keywords = []
                
                # Log raw data for debugging
                logging.info(f"Raw job_stability data for eval {row[0]}: {row[6]}")
                logging.info(f"Raw career_progression data for eval {row[0]}: {row[7]}")
                
                # Handle job stability data
                job_stability_data = row[6]
                if job_stability_data:
                    try:
                        job_stability = safe_json_loads(job_stability_data, {})
                        # Ensure it has the expected structure
                        if not isinstance(job_stability, dict):
                            job_stability = {}
                    except Exception as e:
                        logging.error(f"Error processing job_stability for eval {row[0]}: {str(e)}")
                        job_stability = {}
                else:
                    job_stability = {}
                
                # Handle career progression data
                career_progression_data = row[7]
                if career_progression_data:
                    try:
                        career_progression = safe_json_loads(career_progression_data, {})
                        # Ensure it has the expected structure
                        if not isinstance(career_progression, dict):
                            career_progression = {}
                    except Exception as e:
                        logging.error(f"Error processing career_progression for eval {row[0]}: {str(e)}")
                        career_progression = {}
                else:
                    career_progression = {}
                
                # Handle questions
                technical_questions = safe_json_loads(row[9], [])
                nontechnical_questions = safe_json_loads(row[10], [])
                
                # Ensure profile_summary is a valid string
                profile_summary = str(row[5]) if row[5] is not None else "No summary available"
                
                # Create a default structure for job_stability if empty
                if not job_stability:
                    job_stability = {
                        "StabilityScore": 0,
                        "AverageJobTenure": "N/A",
                        "JobCount": 0,
                        "RiskLevel": "N/A",
                        "ReasoningExplanation": "No job stability data available."
                    }
                
                # Create a default structure for career_progression if empty
                if not career_progression:
                    career_progression = {
                        "progression_score": 0,
                        "key_observations": [],
                        "career_path": [],
                        "red_flags": [],
                        "reasoning": "No career progression data available."
                    }
                
                # Ensure all data is properly serialized for the template
                # This is critical to avoid issues with the template's tojson filter
                try:
                    # Test serialization to catch any issues
                    json.dumps(job_stability)
                    json.dumps(career_progression)
                    json.dumps(technical_questions)
                    json.dumps(nontechnical_questions)
                except (TypeError, ValueError) as e:
                    logging.error(f"Serialization error for evaluation {row[0]}: {str(e)}")
                    # If there's an error, convert to string representation
                    if not isinstance(job_stability, dict):
                        job_stability = {"error": "Invalid data structure", "message": str(job_stability)}
                    if not isinstance(career_progression, dict):
                        career_progression = {"error": "Invalid data structure", "message": str(career_progression)}
                    if not isinstance(technical_questions, list):
                        technical_questions = ["Error loading technical questions"]
                    if not isinstance(nontechnical_questions, list):
                        nontechnical_questions = ["Error loading non-technical questions"]
                
                evaluation = {
                    'id': row[0],
                    'filename': row[1],
                    'job_title': row[2],
                    'match_percentage': row[3],
                    'missing_keywords': missing_keywords,
                    'profile_summary': profile_summary,
                    'job_stability': job_stability,
                    'career_progression': career_progression,
                    'timestamp': row[8],
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'oorwin_job_id': row[11]
                }
                evaluations.append(evaluation)
                
                # Log the processed data for debugging
                logging.info(f"Processed evaluation {row[0]}: job_stability={job_stability}, career_progression={career_progression}")
                
            except Exception as e:
                logging.error(f"Error processing row for evaluation {row[0]}: {str(e)}")
                continue
            
        return render_template('history.html', evaluations=evaluations)
    
    except Exception as e:
        logging.error(f"Error in history route: {str(e)}")
        return render_template('history.html', evaluations=[], error="Failed to load evaluations")
    
    finally:
        conn.close()

@app.route('/feedback_history')
def feedback_history():
    """Display unified feedback history page"""
    return render_template('feedback_history.html')

@app.route('/api/feedback/all')
def get_all_feedback():
    """Get all feedback from all 3 sources"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()

        # Get HR Assistant feedback
        cursor.execute("""
            SELECT qf.id, qf.question_id, qf.rating, qf.feedback, qf.timestamp,
                   qh.question, qh.final_answer
            FROM qa_feedback qf
            JOIN qa_history qh ON qf.question_id = qh.id
            ORDER BY qf.timestamp DESC
        """)
        hr_assistant = []
        for row in cursor.fetchall():
            hr_assistant.append({
                'id': row[0],
                'question_id': row[1],
                'rating': row[2],
                'feedback': row[3],
                'timestamp': row[4],
                'question': row[5],
                'answer': row[6]
            })

        # Get Handbook feedback
        cursor.execute("""
            SELECT hf.id, hf.handbook_id, hf.rating, hf.comments, hf.timestamp,
                   rh.job_title, rh.oorwin_job_id, rh.markdown_content
            FROM handbook_feedback hf
            JOIN recruiter_handbooks rh ON hf.handbook_id = rh.id
            ORDER BY hf.timestamp DESC
        """)
        handbooks = []
        for row in cursor.fetchall():
            handbooks.append({
                'id': row[0],
                'handbook_id': row[1],
                'rating': row[2],
                'comments': row[3],
                'timestamp': row[4],
                'job_title': row[5],
                'oorwin_job_id': row[6],
                'markdown_content': row[7]
            })

        # Get Evaluation feedback
        cursor.execute("""
            SELECT f.id, f.evaluation_id, f.rating, f.comments, f.timestamp,
                   e.filename, e.job_title, e.match_percentage, e.oorwin_job_id,
                   e.match_factors, e.profile_summary, e.missing_keywords,
                   e.job_stability, e.career_progression, e.technical_questions,
                   e.nontechnical_questions, e.behavioral_questions
            FROM feedback f
            JOIN evaluations e ON f.evaluation_id = e.id
            ORDER BY f.timestamp DESC
        """)
        evaluations = []
        for row in cursor.fetchall():
            evaluations.append({
                'id': row[0],
                'evaluation_id': row[1],
                'rating': row[2],
                'comments': row[3],
                'timestamp': row[4],
                'filename': row[5],
                'job_title': row[6],
                'match_percentage': row[7],
                'oorwin_job_id': row[8],
                'match_factors': row[9],
                'profile_summary': row[10],
                'missing_keywords': row[11],
                'job_stability': row[12],
                'career_progression': row[13],
                'technical_questions': row[14],
                'nontechnical_questions': row[15],
                'behavioral_questions': row[16]
            })

        conn.close()

        return jsonify({
            'success': True,
            'hr_assistant': hr_assistant,
            'handbooks': handbooks,
            'evaluations': evaluations
        })

    except Exception as e:
        logging.error(f"Error getting all feedback: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback/check/<feedback_type>/<int:item_id>')
def check_feedback_exists(feedback_type, item_id):
    """Check if feedback already exists for a given item"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        if feedback_type == 'qa':
            cursor.execute("SELECT id FROM qa_feedback WHERE question_id = ?", (item_id,))
        elif feedback_type == 'handbook':
            cursor.execute("SELECT id FROM handbook_feedback WHERE handbook_id = ?", (item_id,))
        elif feedback_type == 'evaluation':
            cursor.execute("SELECT id FROM feedback WHERE evaluation_id = ?", (item_id,))
        else:
            return jsonify({'success': False, 'error': 'Invalid feedback type'}), 400
        
        exists = cursor.fetchone() is not None
        conn.close()
        
        return jsonify({'success': True, 'exists': exists})
        
    except Exception as e:
        logging.error(f"Error checking feedback: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/feedback/handbook', methods=['POST'])
def submit_handbook_feedback():
    """Submit feedback for a recruiter handbook"""
    try:
        data = request.get_json()
        
        if not data or 'handbook_id' not in data or 'rating' not in data:
            return jsonify({'success': False, 'error': 'Missing handbook_id or rating'}), 400
        
        handbook_id = data['handbook_id']
        rating = data['rating']
        comments = data.get('comments', '')
        
        # Validate rating
        if not isinstance(rating, int) or rating < 1 or rating > 5:
            return jsonify({'success': False, 'error': 'Rating must be between 1 and 5'}), 400
        
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Check if feedback already exists
        cursor.execute("SELECT id FROM handbook_feedback WHERE handbook_id = ?", (handbook_id,))
        if cursor.fetchone():
            conn.close()
            return jsonify({'success': False, 'error': 'Feedback already submitted for this handbook'}), 400
        
        # Insert feedback
        cursor.execute("""
            INSERT INTO handbook_feedback (handbook_id, rating, comments)
            VALUES (?, ?, ?)
        """, (handbook_id, rating, comments))
        
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
        
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Feedback already submitted for this handbook'}), 400
    except Exception as e:
        logging.error(f"Error submitting handbook feedback: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# HR Assistant Routes
@app.route('/api/ask', methods=['POST'])
def ask_question():
    try:
        data = request.get_json()
        question = data.get('question')
        online_mode = data.get('online_mode', False)

        if not question:
            return jsonify({'error': 'No question provided'}), 400

        def generate():
            complete_response = []  # Store complete response
            try:
                # Check for special queries first
                special_response = handle_special_queries(question)
                if special_response:
                    complete_response.append(special_response)
                    yield special_response
                    return

                # Expand acronyms in the question
                expanded_question = expand_acronyms(question)

                if online_mode:
                    # ONLINE MODE: Answer any general question using LLM knowledge
                    # No RAG constraints - can answer anything
                    detailed_prompt = f"""You are an expert AI assistant. Provide a comprehensive and detailed answer to the following question. Your response should be thorough, well-structured, and accurate.

                    Question: {expanded_question}

                    Instructions:
                    1. Use your knowledge to provide a complete answer
                    2. If the question is about HR policies, benefits, or company-specific information, note that you may not have the latest company-specific details
                    3. Format your response with clear sections and bullet points where appropriate
                    4. Include relevant examples and context
                    5. If you're uncertain about specific facts, mention that
                    """
                    
                    response = model.generate_content(detailed_prompt, stream=True)
                    for chunk in response:
                        if chunk.text:
                            complete_response.append(chunk.text)
                            yield chunk.text
                else:
                    # RAG MODE: Strict retrieval from local documents only
                    # Use hybrid search (BM25 + Vector) for better coverage
                    
                    # Step 1: Hybrid retrieval - combine BM25 and Vector search
                    all_retrieved_docs = []
                    
                    # Vector search (semantic similarity) - increased k for better coverage
                    if vectorstore is not None:
                        vector_docs = vectorstore.similarity_search(expanded_question, k=15)
                        all_retrieved_docs.extend([(doc, 'vector') for doc in vector_docs])
                        logging.info(f"🔍 Vector search retrieved {len(vector_docs)} documents")
                    
                    # BM25 search (keyword matching) - better for exact terms and tables
                    if bm25_index and bm25_corpus:
                        try:
                            query_tokens = expanded_question.lower().split()
                            bm25_scores = bm25_index.get_scores(query_tokens)
                            
                            # Get top BM25 results - increased k for better coverage
                            top_indices = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)[:10]
                            bm25_results = []
                            for idx in top_indices:
                                if bm25_scores[idx] > 0:  # Only include relevant results
                                    text_content = " ".join(bm25_corpus[idx])
                                    # Create a Document-like object with metadata for consistency
                                    from langchain_core.documents import Document as LangchainDocument
                                    # Get metadata for this chunk if available
                                    metadata = bm25_metadata[idx] if idx < len(bm25_metadata) else {}
                                    bm25_doc = LangchainDocument(page_content=text_content, metadata=metadata)
                                    bm25_results.append(bm25_doc)
                                    all_retrieved_docs.append((bm25_doc, 'bm25'))
                            
                            logging.info(f"🔍 BM25 search retrieved {len(bm25_results)} documents")
                        except Exception as e:
                            logging.warning(f"⚠️ BM25 search failed: {e}")
                    
                    # Step 2: Deduplicate and prioritize
                    seen_content = set()
                    unique_docs = []
                    table_docs = []
                    text_docs = []
                    
                    for doc, source in all_retrieved_docs:
                        content_hash = hash(doc.page_content[:100])  # Hash first 100 chars for dedup
                        if content_hash not in seen_content:
                            seen_content.add(content_hash)
                            # Classify as table or text
                            if "[TABLE DATA]" in doc.page_content or ("|" in doc.page_content and doc.page_content.count("|") > 3):
                                table_docs.append(doc)
                            else:
                                text_docs.append(doc)
                            unique_docs.append(doc)
                    
                    # Step 3: Prioritize tables and limit context window
                    # Tables first (they're often most precise), then text chunks
                    prioritized_docs = table_docs + text_docs
                    context_docs = prioritized_docs[:12]  # Take top 12 for context
                    
                    # Step 4: Build context with proper source citations (filename + page)
                    if context_docs:
                        context_parts = []
                        for i, doc in enumerate(context_docs):
                            # Extract actual source filename and page from metadata
                            source_name = doc.metadata.get('source', 'Unknown Document') if hasattr(doc, 'metadata') and doc.metadata else 'Unknown Document'
                            page_num = doc.metadata.get('page', 'N/A') if hasattr(doc, 'metadata') and doc.metadata else 'N/A'
                            
                            # Format citation with actual filename
                            if source_name != 'Unknown Document':
                                citation = f"{source_name}"
                                if page_num != 'N/A':
                                    citation += f", page {page_num}"
                            else:
                                citation = f"Source {i+1}"
                            
                            # Add relevance indicator for tables with proper citation
                            if doc in table_docs:
                                context_parts.append(f"[RELEVANT TABLE DATA - {citation}]\n{doc.page_content}")
                            else:
                                context_parts.append(f"[RELEVANT CONTEXT - {citation}]\n{doc.page_content}")
                        context = "\n\n---\n\n".join(context_parts)
                        
                        # Log retrieval stats
                        logging.info(f"📚 Total unique documents retrieved: {len(unique_docs)} ({len(table_docs)} tables, {len(text_docs)} text)")
                        if table_docs:
                            logging.info(f"📊 Including {len(table_docs)} table chunks in context")
                    else:
                        context = ""
                        logging.warning("⚠️ No documents retrieved from knowledge base")
                    
                    # Step 5: Enhanced RAG prompt with strict enforcement
                    if context:
                        prompt = f"""You are an expert HR Assistant. Answer the question STRICTLY based on the provided context from company policy documents.

                    Question: {expanded_question}
                    
                    Context from Company Documents:
                    {context}

                    STRICT RULES FOR FORMATTING:
                    
                    1. **ONLY use information from the context provided above.** Do not use external knowledge.
                    
                    2. **TABLE FORMATTING (CRITICAL)**:
                       - If the context contains tables (look for markers like [TABLE DATA] ... [END TABLE]), you MUST reproduce the markdown table **verbatim** from the context
                       - Use this EXACT format:
                         ```
                         | Column 1 | Column 2 | Column 3 |
                         |---------|---------|----------|
                         | Data 1  | Data 2  | Data 3   |
                         | Data 4  | Data 5  | Data 6   |
                         ```
                       - Always include the header separator row (|---------|---------|)
                       - Keep table columns aligned and readable
                       - If a table is complex, break it into smaller, clearer tables
                       - Add a brief title above each table (e.g., "### Performance Rating Table")
                    
                    3. **RESPONSE STRUCTURE**:
                       - Start with a brief 1-2 sentence summary
                       - Use clear headings (## for main sections, ### for subsections)
                       - Use bullet points (• or -) for lists, NOT long paragraphs
                       - Add blank lines between sections for readability
                       - Keep paragraphs SHORT (3-4 sentences max)
                       - Use bold (**text**) for key terms and important points
                    
                    4. **SOURCE CITATION**:
                       - **ALWAYS cite sources using actual document names** (e.g., "According to [Leave Policy.pdf, page 5]")
                       - Place citations at the END of sentences or paragraphs, not mid-sentence
                       - Format: `[Document Name.pdf, page X]`
                       - Use actual filename, not generic "Source 1" or "Source 2"
                    
                    5. **READABILITY ENHANCEMENTS**:
                       - Break dense information into digestible chunks
                       - Use numbered lists (1., 2., 3.) for step-by-step processes
                       - Use bullet lists (•) for features, benefits, or items
                       - Add horizontal rules (---) to separate major sections
                       - Use emojis sparingly for visual breaks (✅, 📋, 📊, etc.)
                    
                    6. **EXAMPLE OF GOOD FORMATTING**:
                       ```
                       ## Performance Appraisal Policy
                       
                       The performance appraisal policy outlines how employees are evaluated annually.
                       
                       ### Key Features
                       • Appraisals are conducted yearly (April to March)
                       • Based on previously agreed KRAs
                       • Results can lead to salary hikes or promotions
                       
                       ### Performance Rating Table
                       | Score | Rating | Description |
                       |-------|--------|-------------|
                       | 5     | Exceptional | Targets met at 200% or above |
                       | 4     | Outstanding | Targets exceeded significantly |
                       | 3     | Good | Consistently met expectations |
                       
                       [APPRAISAL & PROMOTION POLICY.pdf, page 1]
                       ```
                    
                    7. **IMPORTANT**: If the question cannot be answered from the provided context, you MUST respond with:
                       - "I'm sorry, but the information about '[topic]' is not available in our company policy documents."
                       - "💡 **Suggestion**: Please enable the **'Go Online'** toggle and try asking your question again."
                    
                    8. DO NOT make up information or use knowledge outside the provided context.
                    
                    Answer:"""
                    else:
                        # No context found - give helpful message
                        prompt = f"""The question "{expanded_question}" could not be answered from the available company policy documents.

                    Please note:
                    - The information may not be in the current knowledge base
                    - The document may need to be updated or added
                    - You can try rephrasing the question or enabling "Go Online" mode for general information
                    
                    Would you like to:
                    1. Try rephrasing your question
                    2. Enable "Go Online" mode for general information
                    3. Contact HR for company-specific policies not yet in the system"""
                        
                        # Still generate response but with this constraint
                        response = model.generate_content(prompt, stream=True)
                        for chunk in response:
                            if chunk.text:
                                complete_response.append(chunk.text)
                                yield chunk.text
                        return
                    
                    response = model.generate_content(prompt, stream=True)
                    for chunk in response:
                        if chunk.text:
                            complete_response.append(chunk.text)
                            yield chunk.text

                # Store the complete Q&A in history after streaming is done
                final_answer = "".join(complete_response)
                conn = sqlite3.connect('combined_db.db')
                c = conn.cursor()
                c.execute('''INSERT INTO qa_history (question, retrieved_docs, final_answer)
                            VALUES (?, ?, ?)''', (question, None, final_answer))
                conn.commit()
                conn.close()

            except Exception as e:
                error_msg = f"Error: {str(e)}"
                yield error_msg
                # Store error in database
                conn = sqlite3.connect('combined_db.db')
                c = conn.cursor()
                c.execute('''INSERT INTO qa_history (question, final_answer)
                            VALUES (?, ?)''', (question, error_msg))
                conn.commit()
                conn.close()

        return Response(stream_with_context(generate()), mimetype='text/plain')

    except Exception as e:
        print(f"Error in ask_question: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route("/api/update_index", methods=["POST"])
def update_index_api():
    """Manually refresh the Pinecone & BM25 index."""
    try:
        # Rebuild BM25 index
        build_bm25_index(POLICIES_FOLDER)
        
        # Repopulate Pinecone
        populate_pinecone_index()
        
        return jsonify({"message": "Indexes updated successfully"}), 200
    except Exception as e:
        logging.error(f"❌ Index Update Error: {e}", exc_info=True)
        return jsonify({"error": "Failed to update indexes"}), 500

# Resume Evaluator Routes
async def async_gemini_generate(prompt):
    """Async wrapper for Gemini generation with improved JSON handling"""
    try:
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Remove any JSON formatting artifacts
        response_text = response_text.replace('\n', ' ')
        response_text = re.sub(r'\s+', ' ', response_text)
        
        # Remove markdown code block markers if present
        response_text = re.sub(r'^```json\s*', '', response_text)
        response_text = re.sub(r'\s*```$', '', response_text)
        
        # Try to parse as JSON
        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            # Try to extract JSON using regex
            json_match = re.search(r'\{.*\}', response_text)
            if json_match:
                try:
                    return json.loads(json_match.group(0))
                except json.JSONDecodeError:
                    logging.error(f"Failed to parse extracted JSON: {json_match.group(0)}")
                    return get_default_career_analysis()
            else:
                logging.error(f"No JSON found in response: {response_text}")
                return get_default_career_analysis()
                
    except Exception as e:
        logging.error(f"Gemini generation error: {str(e)}")
        return get_default_career_analysis()

async def async_analyze_stability(resume_text):
    """Async job stability analysis"""
    try:
        stability_prompt = job_stability_prompt.format(resume_text=resume_text)
        response = await async_gemini_generate(stability_prompt)
        
        if not response:
            raise ValueError("Failed to get stability analysis")
            
        # Ensure all required fields exist
        default_data = {
            "IsStable": True,
            "AverageJobTenure": "Unknown",
            "JobCount": 0,
            "StabilityScore": 0,
            "ReasoningExplanation": "Could not analyze job stability",
            "RiskLevel": "Unknown"
        }
        
        # Merge response with defaults
        for key, default_value in default_data.items():
            if key not in response:
                response[key] = default_value
                
        return response
        
    except Exception as e:
        logging.error(f"Error in async_analyze_stability: {str(e)}")
        return {
            "IsStable": True,
            "AverageJobTenure": "Unknown",
            "JobCount": 0,
            "StabilityScore": 0,
            "ReasoningExplanation": "Could not analyze job stability",
            "RiskLevel": "Unknown"
        }

async def async_generate_questions(resume_text, job_description, profile_summary):
    """Async interview questions generation"""
    try:
        questions_prompt = interview_questions_prompt.format(
            resume_text=resume_text,
            job_description=job_description,
            profile_summary=profile_summary
        )
        response = await async_gemini_generate(questions_prompt)
        
        if not response:
            raise ValueError("Failed to generate interview questions")
            
        # Ensure we have the required fields with proper defaults
        default_data = {
            "TechnicalQuestions": [],
            "NonTechnicalQuestions": []
        }
        
        # Merge response with defaults
        for key, default_value in default_data.items():
            if key not in response:
                response[key] = default_value
            elif not isinstance(response[key], list):
                response[key] = [str(response[key])] if response[key] else []
                
        return response
        
    except Exception as e:
        logging.error(f"Error in async_generate_questions: {str(e)}")
        return {
            "TechnicalQuestions": [],
            "NonTechnicalQuestions": []
        }

async def async_generate_recruiter_handbook(resume_text, job_description):
    """Async recruiter handbook generation - returns markdown text"""
    try:
        handbook_prompt = recruiter_handbook_prompt.format(
            resume_text=resume_text,
            job_description=job_description
        )
        
        # Use Gemini to generate the recruiter handbook (run in thread pool to avoid blocking)
        # This returns markdown text, not JSON
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, lambda: model.generate_content(handbook_prompt))
        response_text = response.text.strip()
        
        if not response_text:
            raise ValueError("Failed to generate recruiter handbook")
        
        # Return just the markdown content (JSON summary removed as per user request)
        return {
            "markdown_content": response_text,
            "json_summary": None
        }
        
    except Exception as e:
        logging.error(f"Error in async_generate_recruiter_handbook: {str(e)}")
        return {
            "markdown_content": f"## Error\n\nFailed to generate recruiter handbook: {str(e)}",
            "json_summary": None
        }

@app.route('/evaluate', methods=['POST'])
async def evaluate_resume():
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')

        if not job_title or not job_description:
            return jsonify({'error': 'Missing job title or description'}), 400

        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract text from resume
        resume_text = extract_text_from_file(file_path)
        if resume_text is None:
            return jsonify({'error': 'Failed to extract text from file'}), 500

        # Generate evaluation using Gemini with optimized parameters
        formatted_prompt = input_prompt_template.format(resume_text=resume_text, job_description=job_description)
        
        try:
            # Run all analyses concurrently using asyncio.gather
            main_response, stability_data, career_data = await asyncio.gather(
                async_gemini_generate(formatted_prompt),
                async_analyze_stability(resume_text),
                analyze_career_progression(resume_text)  # Now properly awaited
            )
            
            if not main_response:
                raise ValueError("Failed to get main evaluation response")
                
            if not career_data:
                career_data = {
                    "progression_score": 50,
                    "key_observations": ["Failed to analyze career progression"],
                    "career_path": [],
                    "red_flags": ["Analysis error"],
                    "reasoning": "Failed to process career data"
                }
                
        except Exception as e:
            logging.error(f"Error during concurrent analysis: {str(e)}")
            return jsonify({'error': 'Failed to analyze resume'}), 500
        
        # Extract values from main response
        match_percentage_str = main_response.get("JD Match", "0%")
        match_percentage = int(match_percentage_str.strip('%'))
        missing_keywords = main_response.get("MissingKeywords", [])
        profile_summary = main_response.get("Profile Summary", "No summary provided.")
        over_under_qualification = main_response.get("Over/UnderQualification Analysis", "No qualification mismatch concerns detected.")
        match_factors = main_response.get("Match Factors", {})
        candidate_fit_analysis = main_response.get("Candidate Fit Analysis", {})

        # Prepare additional information
        additional_info = {
            "job_stability": stability_data,
            "career_progression": career_data,
            "reasoning": main_response.get("Reasoning", "")
        }

        # Generate unique ID for evaluation
        eval_id = str(uuid.uuid4())

        # Save evaluation to database with additional info
        db_id = save_evaluation(eval_id, filename, job_title, match_percentage, missing_keywords, profile_summary, match_factors, stability_data, additional_info, None, candidate_fit_analysis, over_under_qualification)
        if db_id:
            # Generate interview questions asynchronously
            questions_data = await async_generate_questions(resume_text, job_description, profile_summary)
            
            technical_questions = questions_data.get("TechnicalQuestions", [])
            nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
            behavioral_questions = QUICK_CHECKS

            # Save interview questions with proper JSON encoding (use database ID)
            if save_interview_questions(db_id, 
                                     json.dumps(technical_questions), 
                                     json.dumps(nontechnical_questions), 
                                     json.dumps(behavioral_questions)):
                return jsonify({
                    'id': eval_id,
                    'match_percentage': match_percentage,
                    'match_percentage_str': match_percentage_str,
                    'missing_keywords': missing_keywords,
                    'profile_summary': profile_summary,
                    'over_under_qualification': over_under_qualification,
                    'match_factors': match_factors,
                    'candidate_fit_analysis': candidate_fit_analysis,
                    'job_stability': stability_data,
                    'career_progression': career_data,
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'behavioral_questions': behavioral_questions
                })
            else:
                return jsonify({'error': 'Failed to save interview questions'}), 500
        else:
            return jsonify({'error': 'Failed to save evaluation'}), 500

    except Exception as e:
        logging.error(f"Error in evaluate_resume: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/evaluate-stream', methods=['POST'])
def evaluate_resume_stream():
    """Streaming version of resume evaluation for better UX"""
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')
        oorwin_job_id = request.form.get('oorwin_job_id', '').strip()  # Get JobID from form

        if not job_title or not job_description:
            return jsonify({'error': 'Missing job title or description'}), 400

        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract text from resume
        resume_text = extract_text_from_file(file_path)
        if resume_text is None:
            return jsonify({'error': 'Failed to extract text from file'}), 500

        # Generate unique ID for evaluation
        eval_id = str(uuid.uuid4())

        def generate():
            try:
                # Send initial response
                yield f"data: {json.dumps({'status': 'processing', 'message': 'Analyzing resume...', 'eval_id': eval_id})}\n\n"
                
                # Step 1: Main resume analysis (most important - show results immediately)
                yield f"data: {json.dumps({'status': 'step1', 'message': 'Evaluating resume against job requirements...'})}\n\n"
                formatted_prompt = input_prompt_template.format(resume_text=resume_text, job_description=job_description)
                main_response = asyncio.run(async_gemini_generate(formatted_prompt))
                
                if not main_response:
                    yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to analyze resume'})}\n\n"
                    return
                
                # Extract basic values
                match_percentage_str = main_response.get("JD Match", "0%")
                match_percentage = int(match_percentage_str.strip('%'))
                missing_keywords = main_response.get("MissingKeywords", [])
                profile_summary = main_response.get("Profile Summary", "No summary provided.")
                over_under_qualification = main_response.get("Over/UnderQualification Analysis", "No qualification mismatch concerns detected.")
                match_factors = main_response.get("Match Factors", {})
                candidate_fit_analysis = main_response.get("Candidate Fit Analysis", {})
                
                # Send basic results immediately (user sees this in ~5-8 seconds instead of 20)
                basic_results = {
                    'status': 'basic_results',
                    'id': eval_id,
                    'match_percentage': match_percentage,
                    'match_percentage_str': match_percentage_str,
                    'missing_keywords': missing_keywords,
                    'profile_summary': profile_summary,
                    'over_under_qualification': over_under_qualification,
                    'match_factors': match_factors,
                    'candidate_fit_analysis': candidate_fit_analysis
                }
                yield f"data: {json.dumps(basic_results)}\n\n"
                
                # Step 2: Run additional analyses in parallel
                yield f"data: {json.dumps({'status': 'step2', 'message': 'Analyzing job stability and career progression...'})}\n\n"
                
                stability_data = asyncio.run(async_analyze_stability(resume_text))
                career_data = asyncio.run(analyze_career_progression(resume_text))
                
                if not career_data:
                    career_data = {
                        "progression_score": 50,
                        "key_observations": ["Failed to analyze career progression"],
                        "career_path": [],
                        "red_flags": ["Analysis error"],
                        "reasoning": "Failed to process career data"
                    }
                
                # Send stability and career data
                additional_data = {
                    'status': 'additional_data',
                    'job_stability': stability_data,
                    'career_progression': career_data
                }
                yield f"data: {json.dumps(additional_data)}\n\n"
                
                # Step 3: Generate interview questions
                yield f"data: {json.dumps({'status': 'step3', 'message': 'Generating interview questions...'})}\n\n"
                questions_data = asyncio.run(async_generate_questions(resume_text, job_description, profile_summary))
                
                technical_questions = questions_data.get("TechnicalQuestions", [])
                nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
                behavioral_questions = QUICK_CHECKS
                
                # Send questions
                questions_data_response = {
                    'status': 'questions',
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'behavioral_questions': behavioral_questions
                }
                yield f"data: {json.dumps(questions_data_response)}\n\n"
                
                # Step 4: Save to database
                yield f"data: {json.dumps({'status': 'step4', 'message': 'Saving results...'})}\n\n"
                
                additional_info = {
                    "job_stability": stability_data,
                    "career_progression": career_data,
                    "reasoning": main_response.get("Reasoning", "")
                }
                
                # Debug: Log the data being saved
                logging.info(f"Attempting to save evaluation: eval_id={eval_id}, filename={filename}, job_title={job_title}")
                
                db_id = save_evaluation(eval_id, filename, job_title, match_percentage, missing_keywords, profile_summary, match_factors, stability_data, additional_info, oorwin_job_id, candidate_fit_analysis, over_under_qualification)
                logging.info(f"Save evaluation result: db_id={db_id}")
                
                if db_id:
                    # Use the database ID (integer) for saving interview questions
                    if save_interview_questions(db_id, 
                                             json.dumps(technical_questions), 
                                             json.dumps(nontechnical_questions), 
                                             json.dumps(behavioral_questions)):
                        # Send the database ID back to frontend for feedback submission
                        yield f"data: {json.dumps({'status': 'complete', 'message': 'Analysis complete!', 'db_id': db_id})}\n\n"
                    else:
                        yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to save interview questions'})}\n\n"
                else:
                    yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to save evaluation'})}\n\n"
                    
            except Exception as e:
                logging.error(f"Error in streaming evaluation: {str(e)}")
                yield f"data: {json.dumps({'status': 'error', 'message': str(e)})}\n\n"

        return Response(stream_with_context(generate()), mimetype='text/event-stream')

    except Exception as e:
        logging.error(f"Error in evaluate_resume_stream: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/get_interview_questions/<evaluation_id>', methods=['GET'])
def get_interview_questions(evaluation_id):
    """Get interview questions for a specific evaluation"""
    conn = None
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # First, get the evaluation details to regenerate questions if needed
        cursor.execute(
            """
            SELECT e.resume_path, e.job_title, e.job_description, e.profile_summary 
            FROM evaluations e 
            WHERE e.id = ?
            """,
            (evaluation_id,)
        )
        eval_result = cursor.fetchone()
        
        # Then get existing questions
        cursor.execute(
            "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
            (evaluation_id,)
        )
        result = cursor.fetchone()
        
        # Initialize default values
        technical_questions = []
        nontechnical_questions = []
        behavioral_questions = []
        
        if result:
            try:
                # Parse saved questions with proper error handling
                def parse_json_safely(json_str):
                    if not json_str:
                        return []
                    try:
                        data = json.loads(json_str)
                        if isinstance(data, list):
                            return data
                        elif isinstance(data, str):
                            try:
                                return json.loads(data)
                            except:
                                return [data]
                        else:
                            return [str(data)]
                    except json.JSONDecodeError:
                        try:
                            # Try to clean and parse the string
                            cleaned_str = json_str.strip('[]"\' ').replace('\\', '')
                            items = [item.strip('"\' ') for item in cleaned_str.split(',')]
                            return [item for item in items if item]
                        except:
                            return []

                technical_questions = parse_json_safely(result[0])
                nontechnical_questions = parse_json_safely(result[1])
                behavioral_questions = parse_json_safely(result[2])
            except Exception as e:
                logging.error(f"Error parsing interview questions: {str(e)}")
                technical_questions = []
                nontechnical_questions = []
                behavioral_questions = []

        # Only regenerate questions if they are completely missing (not just empty)
        # This prevents regenerating questions when they exist but are empty arrays
        if not result and eval_result:
            logging.info(f"No interview questions found in database for evaluation {evaluation_id}, generating new ones")
            resume_text = extract_text_from_file(eval_result[0])
            if resume_text:
                        questions_data = asyncio.run(async_generate_questions(
                            resume_text,
                            eval_result[2],  # job_description
                            eval_result[3]   # profile_summary
                        ))
                        
            technical_questions = questions_data.get("TechnicalQuestions", [])
            nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
            behavioral_questions = QUICK_CHECKS
                        
             # Save regenerated questions
            cursor.execute(
                            """
                    INSERT INTO interview_questions 
                    (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (evaluation_id,
                     json.dumps(technical_questions), 
                             json.dumps(nontechnical_questions), 
                     json.dumps(behavioral_questions))
                        )
            conn.commit()
            logging.info(f"Generated and saved new questions for evaluation {evaluation_id}")

            return jsonify({
                    'technical_questions': technical_questions or ["No technical questions available"],
                    'nontechnical_questions': nontechnical_questions or ["No non-technical questions available"],
                    'behavioral_questions': behavioral_questions or QUICK_CHECKS
                })

    except Exception as e:
        logging.error(f"Database error in get_interview_questions: {str(e)}")
        return jsonify({
            'technical_questions': ["Error loading technical questions"],
            'nontechnical_questions': ["Error loading non-technical questions"],
            'behavioral_questions': QUICK_CHECKS
        })
    finally:
        if conn:
            conn.close()

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """Handle feedback for both Q&A and resume evaluations."""
    try:
        data = request.get_json()
        logging.info(f"Received feedback data: {data}")
        
        if not data:
            logging.error("No feedback data received")
            return jsonify({'error': 'No feedback data provided'}), 400
            
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        try:
            # Check if this is Q&A feedback
            if 'question' in data:
                if 'rating' not in data:
                    return jsonify({'error': 'Missing rating'}), 400
                
                # Get question_id from qa_history
                cursor.execute("""
                    SELECT id FROM qa_history 
                    WHERE question = ? 
                    ORDER BY timestamp DESC 
                    LIMIT 1
                """, (data['question'],))
                
                result = cursor.fetchone()
                if not result:
                    # If question not found, create a new entry
                    cursor.execute("""
                        INSERT INTO qa_history (question, final_answer)
                        VALUES (?, ?)
                    """, (data['question'], ''))
                    question_id = cursor.lastrowid
                else:
                    question_id = result[0]
                
                # Check if feedback already exists for this question
                cursor.execute("SELECT id FROM qa_feedback WHERE question_id = ?", (question_id,))
                if cursor.fetchone():
                    return jsonify({'error': 'Feedback already submitted for this question'}), 400
                
                # Insert feedback
                cursor.execute("""
                    INSERT INTO qa_feedback (question_id, rating, feedback, timestamp)
                    VALUES (?, ?, ?, datetime('now'))
                """, (question_id, data['rating'], data.get('feedback', '')))
                
            else:
                # Handle resume evaluation feedback
                if 'evaluation_id' not in data or 'rating' not in data:
                    return jsonify({'error': 'Missing evaluation_id or rating'}), 400
                
                # Check if feedback already exists
                cursor.execute("SELECT id FROM feedback WHERE evaluation_id = ?", (data['evaluation_id'],))
                if cursor.fetchone():
                    return jsonify({'error': 'Feedback already submitted for this evaluation'}), 400
                
                # Insert feedback into the feedback table
                cursor.execute("""
                    INSERT INTO feedback (evaluation_id, rating, comments, timestamp)
                    VALUES (?, ?, ?, datetime('now'))
                """, (data['evaluation_id'], data['rating'], data.get('comments', '')))
            
            conn.commit()
            return jsonify({'message': 'Feedback submitted successfully'})
            
        finally:
            conn.close()
            
    except sqlite3.IntegrityError as e:
        logging.error(f"Integrity error in submit_feedback: {str(e)}")
        return jsonify({'error': 'Feedback already submitted'}), 400
    except sqlite3.Error as e:
        logging.error(f"Database error in submit_feedback: {str(e)}")
        return jsonify({'error': 'Database error occurred'}), 500
    except Exception as e:
        logging.error(f"Error in submit_feedback: {str(e)}")
        return jsonify({'error': 'An unexpected error occurred'}), 500

# --- Document Processing ---
# Optimized chunking for HR policy documents:
# - Larger chunks (1200) preserve context and complete policy explanations
# - Higher overlap (250) ensures continuity across chunks
# - Better separators prioritize paragraph/sentence boundaries
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1200,        # Increased from 400 for better context retention
    chunk_overlap=250,      # Increased from 50 (20% overlap for continuity)
    separators=["\n\n", "\n", ". ", " ", ""]  # Better paragraph/sentence awareness
)

def dataframe_to_clean_markdown(df: pd.DataFrame) -> str:
    """Convert a pandas DataFrame into a clean GitHub-flavored markdown table.

    - Fills missing/None headers with generic names (Column 1, Column 2, ...)
    - Collapses multi-line cell content into single lines
    - Ensures string-typed cells
    - Removes the index column from the markdown output
    """
    try:
        # Defensive copy
        table_df = df.copy()

        # Normalize headers
        normalized_columns = []
        for i, col in enumerate(table_df.columns):
            col_str = str(col).strip() if col is not None else ""
            if col_str == "" or col_str.lower() == "none":
                col_str = f"Column {i+1}"
            normalized_columns.append(col_str)
        table_df.columns = normalized_columns

        # Drop index-like first column if it looks like 0..N sequence
        if table_df.shape[1] > 0:
            first_col_vals = list(table_df.iloc[:, 0])
            def looks_like_sequential_index(vals):
                try:
                    ints = [int(str(v)) for v in vals]
                    return ints == list(range(len(ints)))
                except Exception:
                    return False
            if looks_like_sequential_index(first_col_vals):
                table_df = table_df.iloc[:, 1:]

        # Drop columns that are entirely empty after stripping
        if table_df.shape[1] > 0:
            non_empty_cols = []
            for c in table_df.columns:
                col_str = table_df[c].astype(str).str.strip()
                if (col_str != "").any():
                    non_empty_cols.append(c)
            if non_empty_cols:
                table_df = table_df[non_empty_cols]

        # Normalize cell content: string type, collapse newlines/tabs, trim
        for c in table_df.columns:
            table_df[c] = (
                table_df[c]
                .astype(str)
                .str.replace("\r\n|\r|\n", " ", regex=True)
                .str.replace("\t", " ", regex=True)
                .str.replace("\\s+", " ", regex=True)
                .str.strip()
            )

        # Render as markdown without index
        return table_df.to_markdown(index=False, tablefmt="pipe")
    except Exception as e:
        logging.warning(f"Failed to render markdown table cleanly, falling back: {e}")
        # Fallback to basic to_markdown if anything goes wrong
        try:
            return df.to_markdown(index=False, tablefmt="pipe")
        except Exception:
            return df.to_markdown()

def process_pdf(pdf_path, documents, table_chunks):
    """Extract text and tables from a PDF file."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # Extract and process text
                text = page.extract_text() or ""
                if text:
                    text_chunks = text_splitter.split_text(text)
                    documents.extend(text_chunks)
                
                # Extract and process tables
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:  # Ensure table has headers and data
                        df = pd.DataFrame(table[1:], columns=table[0])
                        table_markdown = dataframe_to_clean_markdown(df)
                        
                        # Enrich table for better BM25 matching (same as Pinecone indexing)
                        column_names = " ".join([str(col).lower() for col in df.columns if col])
                        sample_values = []
                        # Use positional access to avoid issues with non-standard/duplicate column labels
                        num_sample_cols = min(3, len(df.columns))
                        for col_idx in range(num_sample_cols):
                            series = df.iloc[:, col_idx]
                            sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                        sample_context = " ".join(sample_values[:10])
                        enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                        table_chunks.append(enriched_table)
    except Exception as e:
        logging.error(f"❌ Error processing PDF {pdf_path}: {e}")

def populate_pinecone_index():
    """Extract content from PDF documents and populate Pinecone index."""
    try:
        documents = []
        table_chunks = []
        texts = []
        metadatas = []
        
        # Get all PDF files from the policies folder
        if not os.path.exists(POLICIES_FOLDER):
            logging.warning(f"Policies folder {POLICIES_FOLDER} does not exist")
            return
            
        pdf_files = [f for f in os.listdir(POLICIES_FOLDER) if f.endswith('.pdf')]
        if not pdf_files:
            logging.warning(f"No PDF files found in {POLICIES_FOLDER}")
            return
            
        total_files = len(pdf_files)
        logging.info(f"📚 Processing {total_files} PDF files")
        
        for idx, filename in enumerate(pdf_files, 1):
            pdf_path = os.path.join(POLICIES_FOLDER, filename)
            logging.info(f"📄 Processing file {idx}/{total_files}: {filename}")
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text:
                        chunks = text_splitter.split_text(text)
                        documents.extend(chunks)
                        for chunk in chunks:
                            texts.append(chunk)
                            metadatas.append({
                                "source": filename,
                                "page": page_num,
                                "type": "text"
                            })
                        logging.info(f"   Page {page_num}: Added {len(chunks)} text chunks")
                    
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table and len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_markdown = dataframe_to_clean_markdown(df)
                            
                            # Improve table representation for better retrieval:
                            # 1. Extract column names as descriptive keywords
                            column_names = " ".join([str(col).lower() for col in df.columns if col])
                            
                            # 2. Extract sample data values as context
                            sample_values = []
                            # Use positional access to avoid issues with non-standard/duplicate column labels
                            num_sample_cols = min(3, len(df.columns))
                            for col_idx in range(num_sample_cols):  # First up to 3 columns
                                series = df.iloc[:, col_idx]
                                sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                            sample_context = " ".join(sample_values[:10])
                            
                            # 3. Create enriched table chunk with context
                            enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                            table_chunks.append(enriched_table)
                            texts.append(enriched_table)
                            metadatas.append({
                                "source": filename,
                                "page": page_num,
                                "type": "table"
                            })
                            logging.info(f"   Page {page_num}: Added table {table_num} (enriched with context)")
    except Exception as e:
        logging.error(f"❌ Error in document processing: {str(e)}")
        raise

    try:
        all_chunks = documents + table_chunks
        total_chunks = len(texts)
        
        if total_chunks == 0:
            raise ValueError("No content extracted from documents")
        
        logging.info(f"📊 Preparing to insert {total_chunks} chunks into Pinecone")
        
        # Initialize Pinecone components
        pc = Pinecone(api_key=PINECONE_API_KEY)
        index = pc.Index(PINECONE_INDEX_NAME)
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
        # Insert in batches
        batch_size = 50  # Reduced batch size for better reliability
        for i in range(0, total_chunks, batch_size):
            batch_texts = texts[i:i + batch_size]
            batch_metas = metadatas[i:i + batch_size]
            PineconeVectorStore.from_texts(
                texts=batch_texts,
                embedding=embeddings,
                index_name=PINECONE_INDEX_NAME,
                metadatas=batch_metas
            )
            logging.info(f"✅ Inserted batch {i//batch_size + 1}/{(total_chunks-1)//batch_size + 1}")
        
        # Verify insertion
        stats = index.describe_index_stats()
        vector_count = stats['total_vector_count']
        logging.info(f"🎉 Successfully populated index with {vector_count} vectors")
        
    except Exception as e:
        logging.error(f"❌ Error in Pinecone operations: {str(e)}")
        raise
#         total_files = len(pdf_files)
#         logging.info(f"📚 Processing {total_files} PDF files")
        
#         for idx, filename in enumerate(pdf_files, 1):
#             pdf_path = os.path.join(POLICIES_FOLDER, filename)
#             logging.info(f"📄 Processing file {idx}/{total_files}: {filename}")
            
#             with pdfplumber.open(pdf_path) as pdf:
#                 for page_num, page in enumerate(pdf.pages, 1):
#                     text = page.extract_text() or ""
#                     if text:
#                         chunks = text_splitter.split_text(text)
#                         documents.extend(chunks)
#                         logging.info(f"   Page {page_num}: Added {len(chunks)} text chunks")
#                     
#                     tables = page.extract_tables()
#                     for table_num, table in enumerate(tables, 1):
#                         if table and len(table) > 1:
#                             df = pd.DataFrame(table[1:], columns=table[0])
#                             table_chunks.append(df.to_markdown())
#                             logging.info(f"   Page {page_num}: Added table {table_num}")
#     except Exception as e:
#         logging.error(f"❌ Error in document processing: {str(e)}")
#         raise

#     try:
#     all_chunks = documents + table_chunks
#     total_chunks = len(all_chunks)
        
#     if total_chunks == 0:
#             raise ValueError("No content extracted from documents")
        
#     logging.info(f"📊 Preparing to insert {total_chunks} chunks into Pinecone")
        
#         # Initialize Pinecone components
#     pc = Pinecone(api_key=PINECONE_API_KEY)
#     index = pc.Index(PINECONE_INDEX_NAME)
#     embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
#     # Insert in batches
#     batch_size = 50  # Reduced batch size for better reliability
#     for i in range(0, total_chunks, batch_size):
#             batch = all_chunks[i:i + batch_size]
#             PineconeVectorStore.from_texts(
#                 texts=batch,
#                 embedding=embeddings,
#                 index_name=PINECONE_INDEX_NAME
#             )
#             logging.info(f"✅ Inserted batch {i//batch_size + 1}/{(total_chunks-1)//batch_size + 1}")
        
#         # Verify insertion
#             stats = index.describe_index_stats()
#             vector_count = stats['total_vector_count']
#             logging.info(f"🎉 Successfully populated index with {vector_count} vectors")
        
#     except Exception as e:
#         logging.error(f"❌ Error in Pinecone operations: {str(e)}")
#         raise

def initialize_pinecone():
    """Initialize Pinecone. Create and populate index if it doesn't exist."""
    try:
        logging.info("🔧 Initializing Pinecone...")
        
        # Check if index exists
        pc = Pinecone(api_key=PINECONE_API_KEY)
        index_name = PINECONE_INDEX_NAME
        
        if index_name in pc.list_indexes().names():
            logging.info(f"📋 Index '{index_name}' already exists - using existing index")
            # Check if index has data
            index = pc.Index(index_name)
            stats = index.describe_index_stats()
            vector_count = stats['total_vector_count']
            
            if vector_count > 0:
                logging.info(f"✅ Index '{index_name}' has {vector_count} vectors - no need to populate")
                return True
            else:
                logging.info(f"⚠️ Index '{index_name}' exists but is empty - populating...")
                populate_pinecone_index()
                return True
        else:
            # Create new index only if it doesn't exist
            logging.info(f"🆕 Index '{index_name}' doesn't exist - creating new index...")
            pc.create_index(
                name=index_name,
                dimension=384,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
            
            # Wait for index to be ready
            time.sleep(10)
            logging.info(f"✅ Index '{index_name}' created successfully")
            
            # Populate the new index
            logging.info("📚 Populating new Pinecone index...")
            populate_pinecone_index()
            
            return True
        
    except Exception as e:
        logging.error(f"❌ Error initializing Pinecone: {str(e)}")
        return False

# --- BM25 Setup ---
bm25_index = None
bm25_corpus = None
bm25_metadata = []  # Store metadata for each BM25 chunk (filename, page, type)

def build_bm25_index(folder_path):
    """Builds BM25 index from policy documents with metadata tracking."""
    global bm25_index, bm25_corpus, bm25_metadata
    
    all_texts = []
    table_chunks = []
    text_metadata = []  # Track metadata for text chunks
    table_metadata = []  # Track metadata for table chunks
    
    # Process all PDF files with metadata
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(folder_path, filename)
            
            # Process PDF with metadata tracking
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text() or ""
                    if text:
                        text_chunks = text_splitter.split_text(text)
                        all_texts.extend(text_chunks)
                        # Add metadata for each text chunk
                        for chunk in text_chunks:
                            text_metadata.append({
                                "source": filename,
                                "page": page_num,
                                "type": "text"
                            })
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table and len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_markdown = dataframe_to_clean_markdown(df)
                            
                            # Enrich table for better BM25 matching
                            column_names = " ".join([str(col).lower() for col in df.columns if col])
                            sample_values = []
                            num_sample_cols = min(3, len(df.columns))
                            for col_idx in range(num_sample_cols):
                                series = df.iloc[:, col_idx]
                                sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                            sample_context = " ".join(sample_values[:10])
                            enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                            table_chunks.append(enriched_table)
                            # Add metadata for table chunk
                            table_metadata.append({
                                "source": filename,
                                "page": page_num,
                                "type": "table"
                            })
    
    # Combine text and tables for indexing
    all_chunks = all_texts + table_chunks
    bm25_metadata = text_metadata + table_metadata
    
    if all_chunks:
        # Tokenize for BM25
        bm25_corpus = [text.split() for text in all_chunks]
        bm25_index = BM25Okapi(bm25_corpus)
        logging.info(f"✅ BM25 index built with {len(bm25_corpus)} document chunks (with metadata tracking)")
    else:
        logging.warning("⚠️ No content found for BM25 indexing")

def expand_query_with_llm(question, llm):
    """Expands user query using LLM to include synonyms but retains original meaning."""
    expansion_prompt = f"""
    Provide alternative phrasings and related terms for: '{question}', 
    ensuring the original word is always included. Include HR-specific terms if applicable.
    """
    try:
        expanded_query = llm.invoke(expansion_prompt).content
        logging.info(f"🔍 Query Expansion: {expanded_query}")
        return expanded_query
    except Exception as e:
        logging.error(f"❌ Query Expansion Failed: {e}")
        return question  # Fall back to the original question

def hybrid_search(question, llm, retriever):
    """Performs hybrid retrieval using BM25 and Pinecone vectors."""
    global bm25_index, bm25_corpus
    
    # Expand query
    expanded_query = expand_query_with_llm(question, llm)
    
    results = []
    
    # Step 1: BM25 Keyword Search
    if bm25_index and bm25_corpus:
        bm25_results = bm25_index.get_top_n(expanded_query.split(), bm25_corpus, n=5)
        bm25_texts = [" ".join(text) for text in bm25_results]
        results.extend(bm25_texts)
        logging.info(f"🔍 BM25 Retrieved {len(bm25_texts)} results")
    
    # Step 2: Vector Search
    pinecone_results = retriever.invoke(expanded_query)
    pinecone_texts = [doc.page_content for doc in pinecone_results]
    results.extend(pinecone_texts)
    
    # Prioritize table content (tables contain | character in markdown)
    table_texts = [text for text in results if "|" in text]
    non_table_texts = [text for text in results if "|" not in text]
    
    # Combine results: tables first, then other content
    combined_results = table_texts + non_table_texts
    
    # Remove duplicates while preserving order
    unique_results = []
    seen = set()
    for text in combined_results:
        # Use a hash of the text as a unique identifier
        text_hash = hash(text)
        if text_hash not in seen:
            seen.add(text_hash)
            unique_results.append(text)
    
    # Join and truncate to avoid token limits
    final_text = "\n\n".join(unique_results)[:5000]
    
    return final_text

def save_qa_to_db(question, retrieved_docs, final_answer, feedback=None):
    """Stores a Q&A pair in SQLite with optional feedback."""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        logging.info(f"Saving Q&A to DB - Question: {question[:50]}...")  # Debug log
        
        query = """
        INSERT INTO qa_history (question, retrieved_docs, final_answer, feedback) 
        VALUES (?, ?, ?, ?)
        """
        cursor.execute(query, (question, retrieved_docs, final_answer, feedback))
        conn.commit()
        
        question_id = cursor.lastrowid
        logging.info(f"✅ Q&A stored successfully with ID: {question_id}")
        return question_id
    except Exception as e:
        logging.error(f"❌ Error saving Q&A to DB: {e}", exc_info=True)
        return None
    finally:
        conn.close()

def setup_llm_chain():
    """Initialize the LLM and retrieval chain."""
    # Initialize LLM with optimized parameters
    llm = ChatGroq(
        # model_name= "mixtral-8x7b-32768", 
        #  model_name= "llama-3.1-8b-instant",
        model_name =  "qwen/qwen3-32b",     #"qwen-2.5-32b",
        groq_api_key=GROQ_API_KEY,
        temperature=0.377,
        max_tokens=32768,
        top_p=0.95,
        presence_penalty=0.1,
        frequency_penalty=0.1,
        streaming=True
    )
    
    # Initialize retriever only if vectorstore is available
    retriever = None
    if vectorstore is not None:
        try:
            retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
            logging.info("✅ Retriever initialized successfully")
        except Exception as e:
            logging.error(f"❌ Error initializing retriever: {e}")
            retriever = None
    else:
        logging.warning("⚠️ Vectorstore not available, retriever will be None")
    
    return llm, None, retriever  # Return llm, qa_chain (None for now), retriever

def expand_acronyms(question):
    """Expand HR-related acronyms in the question."""
    expanded_question = question.lower()
    for acronym, full_form in ACRONYM_MAP.items():
        expanded_question = expanded_question.replace(acronym.lower(), full_form.lower())
    return expanded_question

async def analyze_career_progression(resume_text):
    """Analyze career progression from resume text using Gemini."""
    try:
        formatted_prompt = f"""You are an expert HR analyst. Analyze this candidate's career progression.
Return ONLY a JSON object with the following structure, no other text:
{{
    "progression_score": <number 0-100>,
    "key_observations": [<list of string observations>],
    "career_path": [
        {{
            "title": "<job title>",
            "company": "<company name>",
            "duration": "<time period>",
            "level": "<Entry/Mid/Senior/Lead/Manager>",
            "progression": "<Promotion/Lateral/Step Back>"
        }}
    ],
    "red_flags": [<list of string concerns>],
    "reasoning": "<analysis explanation>"
}}

Resume text:
{resume_text}"""

        # Get response from Gemini
        response = await async_gemini_generate(formatted_prompt)
        
        # If response is already a dict (from async_gemini_generate)
        if isinstance(response, dict):
            parsed_response = response
        else:
            try:
                parsed_response = json.loads(response) if isinstance(response, str) else {}
            except json.JSONDecodeError:
                logging.error(f"Failed to parse response as JSON: {response}")
                return get_default_career_analysis()

        # Validate and clean the response data
        cleaned_data = {
            "progression_score": validate_progression_score(parsed_response.get("progression_score", 50)),
            "key_observations": validate_list(parsed_response.get("key_observations", [])) or ["No key observations found"],
            "career_path": validate_career_path(parsed_response.get("career_path", [])),
            "red_flags": validate_list(parsed_response.get("red_flags", [])) or ["No red flags identified"],
            "reasoning": str(parsed_response.get("reasoning", "No analysis provided")).strip()
        }

        # Ensure we have valid data
        if cleaned_data["progression_score"] == 50 and not cleaned_data["career_path"]:
            return get_default_career_analysis()
            
        return cleaned_data

    except Exception as e:
        logging.error(f"Career progression analysis error: {str(e)}")
        logging.error(f"Full traceback:", exc_info=True)
        return get_default_career_analysis()

def get_default_career_analysis():
    """Return default career analysis structure"""
    return {
        "progression_score": 50,
        "key_observations": ["Unable to analyze career progression"],
        "career_path": [],
        "red_flags": ["Analysis encountered technical issues"],
        "reasoning": "Analysis failed to complete"
    }

def validate_progression_score(score):
    """Validate and normalize progression score"""
    try:
        if isinstance(score, str):
            score = score.strip('%')
        score = float(score)
        return int(max(0, min(100, score)))
    except (ValueError, TypeError):
        return 50

def validate_list(items):
    """Validate and clean list items"""
    if not isinstance(items, list):
        return []
    return [str(item).strip() for item in items if item and str(item).strip()]

def validate_career_path(path):
    """Validate and clean career path entries"""
    if not isinstance(path, list):
        return []
    
    cleaned_path = []
    required_fields = ["title", "company", "duration", "level", "progression"]
    
    for entry in path:
        if not isinstance(entry, dict):
            continue
        
        cleaned_entry = {}
        for field in required_fields:
            cleaned_entry[field] = str(entry.get(field, "Not specified")).strip()
        cleaned_path.append(cleaned_entry)
    
    return cleaned_path

def update_db_schema():
    """Update database schema if needed"""
    conn = sqlite3.connect('combined_db.db')
    cursor = conn.cursor()
    
    # Add new columns if they don't exist
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN job_stability TEXT;
        ''')
    except sqlite3.OperationalError:
        pass  # Column already exists
        
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN career_progression TEXT;
        ''')
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Add oorwin_job_id column to evaluations table
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN oorwin_job_id TEXT;
        ''')
        logging.info("Added oorwin_job_id column to evaluations table")
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Create recruiter_handbooks table
    try:
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS recruiter_handbooks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                oorwin_job_id TEXT,
                job_title TEXT,
                job_description TEXT,
                additional_context TEXT,
                markdown_content TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        logging.info("Created recruiter_handbooks table")
    except sqlite3.OperationalError:
        pass  # Table already exists
    
    # Create index on oorwin_job_id for faster queries
    try:
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_evaluations_job_id 
            ON evaluations(oorwin_job_id)
        ''')
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_handbooks_job_id 
            ON recruiter_handbooks(oorwin_job_id)
        ''')
        logging.info("Created indexes on oorwin_job_id")
    except sqlite3.OperationalError:
        pass  # Index already exists
    
    conn.commit()
    conn.close()

@app.route('/api/evaluation/<evaluation_id>', methods=['GET'])
def get_evaluation_details(evaluation_id):
    """API endpoint to get evaluation details by ID"""
    conn = None
    try:
        logging.info(f"Fetching evaluation details for ID: {evaluation_id}")
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Helper function for parsing JSON safely
        def parse_json_safely(json_str):
            if not json_str:
                logging.info("Empty JSON string, returning empty list")
                return []
            try:
                data = json.loads(json_str)
                if isinstance(data, list):
                    logging.info(f"Successfully parsed list with {len(data)} items")
                    return data
                elif isinstance(data, str):
                    try:
                        parsed_data = json.loads(data)
                        logging.info(f"Successfully parsed nested JSON string")
                        return parsed_data
                    except:
                        logging.info(f"Failed to parse nested JSON, treating as single item")
                        return [data]
                else:
                    logging.info(f"Non-list data type: {type(data)}, converting to string")
                    return [str(data)]
            except json.JSONDecodeError as e:
                logging.warning(f"JSON decode error: {str(e)}, attempting cleanup")
                try:
                    # Try to clean and parse the string
                    cleaned_str = json_str.strip('[]"\' ').replace('\\', '')
                    items = [item.strip('"\' ') for item in cleaned_str.split(',')]
                    result = [item for item in items if item]
                    logging.info(f"Cleanup successful, extracted {len(result)} items")
                    return result
                except Exception as e2:
                    logging.error(f"Cleanup failed: {str(e2)}")
                    return []
        
        # Get evaluation details first to get job title for default questions
        cursor.execute('''
            SELECT 
                e.id, 
                e.filename, 
                e.job_title, 
                e.match_percentage, 
                e.profile_summary, 
                e.job_stability,
                e.career_progression,
                e.timestamp,
                e.missing_keywords,
                e.behavioral_questions,
                e.technical_questions,
                e.nontechnical_questions
            FROM evaluations e
            WHERE e.id = ?
        ''', (evaluation_id,))
        
        row = cursor.fetchone()
        if not row:
            logging.warning(f"No evaluation found with ID: {evaluation_id}")
            return jsonify({'error': 'Evaluation not found'}), 404
        
        logging.info(f"Found evaluation with ID: {row[0]}, filename: {row[1]}")
        job_title = row[2]
        
        # Parse JSON fields
        try:
            job_stability = json.loads(row[5]) if row[5] else {}
            logging.info(f"Parsed job_stability: {type(job_stability)}")
        except Exception as e:
            logging.error(f"Error parsing job_stability: {str(e)}")
            job_stability = {}
            
        try:
            career_progression = json.loads(row[6]) if row[6] else {}
            logging.info(f"Parsed career_progression: {type(career_progression)}")
        except Exception as e:
            logging.error(f"Error parsing career_progression: {str(e)}")
            career_progression = {}
        
        # Parse missing keywords with special handling
        try:
            missing_keywords_raw = row[8]
            if missing_keywords_raw:
                try:
                    missing_keywords = json.loads(missing_keywords_raw)
                    logging.info(f"Parsed missing_keywords: {type(missing_keywords)}")
                    # If it's not a list, try to convert it
                    if not isinstance(missing_keywords, list):
                        if isinstance(missing_keywords, str):
                            # Remove brackets and split by commas
                            missing_keywords = [k.strip(' "\'') for k in missing_keywords.strip('[]').split(',')]
                        else:
                            missing_keywords = [str(missing_keywords)]
                except Exception as e:
                    logging.error(f"Error parsing missing_keywords JSON: {str(e)}")
                    # If JSON parsing fails, try to extract from string
                    if isinstance(missing_keywords_raw, str):
                        # Check if it looks like a list
                        if missing_keywords_raw.startswith('[') and missing_keywords_raw.endswith(']'):
                            # Remove brackets and split by commas
                            missing_keywords = [k.strip(' "\'') for k in missing_keywords_raw.strip('[]').split(',')]
                        else:
                            missing_keywords = [missing_keywords_raw]
                    else:
                        missing_keywords = []
            else:
                missing_keywords = []
        except Exception as e:
            logging.error(f"Error processing missing_keywords: {str(e)}")
            missing_keywords = []
        
        # Initialize question variables
        technical_questions = []
        nontechnical_questions = []
        behavioral_questions = []
        
        # PRIORITY 1: Try to get interview questions from interview_questions table FIRST
        # This is the dedicated table for storing interview questions
        cursor.execute(
            "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
            (evaluation_id,)
        )
        iq_result = cursor.fetchone()
        
        if not iq_result:
            logging.info(f"No interview questions found with numeric ID, trying string ID")
            # If no results, try with the string representation of the ID
            cursor.execute(
                "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
                (str(evaluation_id),)
            )
            iq_result = cursor.fetchone()
        
        if iq_result:
            logging.info(f"Found interview questions in interview_questions table for evaluation {evaluation_id}")
            technical_questions = parse_json_safely(iq_result[0])
            nontechnical_questions = parse_json_safely(iq_result[1])
            behavioral_questions = parse_json_safely(iq_result[2])
            logging.info(f"Retrieved from interview_questions table: {len(technical_questions)} technical, {len(nontechnical_questions)} non-technical, {len(behavioral_questions)} behavioral questions")
        else:
            logging.info(f"No interview questions found in interview_questions table for evaluation {evaluation_id}")
            
            # PRIORITY 2: Fallback to evaluations table if interview_questions table is empty
            logging.info("Falling back to evaluations table for questions")
        
        # Try to get behavioral questions from evaluations
        try:
            behavioral_questions_raw = row[9]
            if behavioral_questions_raw:
                behavioral_questions = parse_json_safely(behavioral_questions_raw)
                logging.info(f"Parsed behavioral_questions from evaluations: {len(behavioral_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing behavioral_questions from evaluations: {str(e)}")
        
        # Try to get technical questions from evaluations
        try:
            if row[10]:
                technical_questions = parse_json_safely(row[10])
                logging.info(f"Parsed technical_questions from evaluations: {len(technical_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing technical_questions from evaluations: {str(e)}")
        
        # Try to get non-technical questions from evaluations
        try:
            if row[11]:
                nontechnical_questions = parse_json_safely(row[11])
                logging.info(f"Parsed nontechnical_questions from evaluations: {len(nontechnical_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing nontechnical_questions from evaluations: {str(e)}")
        
        # If still no behavioral questions, use default QUICK_CHECKS
        if not behavioral_questions:
            logging.info("No behavioral questions found, using QUICK_CHECKS")
            behavioral_questions = QUICK_CHECKS
        
        
        # Only generate default questions if we still don't have any questions at all
        if not technical_questions and not nontechnical_questions:
            logging.info(f"Generating default questions for job title: {job_title}")
            default_technical, default_nontechnical = get_default_interview_questions(job_title)
            
            if not technical_questions:
                technical_questions = default_technical
                logging.info(f"Using default technical questions: {len(technical_questions)} questions")
            
            if not nontechnical_questions:
                nontechnical_questions = default_nontechnical
                logging.info(f"Using default non-technical questions: {len(nontechnical_questions)} questions")
        
        # Create response
        response = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'match_percentage': row[3],
            'profile_summary': row[4] or "No summary available",
            'job_stability': job_stability,
            'career_progression': career_progression,
            'timestamp': row[7],
            'missing_keywords': missing_keywords,
            'technical_questions': technical_questions,
            'nontechnical_questions': nontechnical_questions,
            'behavioral_questions': behavioral_questions
        }
        
        logging.info(f"Returning response with {len(technical_questions)} technical questions, {len(nontechnical_questions)} non-technical questions, {len(behavioral_questions)} behavioral questions")
        return jsonify(response)
    
    except Exception as e:
        logging.error(f"Error fetching evaluation details: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
    finally:
        if conn:
            conn.close()

@app.route('/api/generate_questions/<evaluation_id>', methods=['POST'])
async def generate_questions_api(evaluation_id):
    """API endpoint to generate interview questions for an evaluation"""
    conn = None
    try:
        logging.info(f"Generating questions for evaluation ID: {evaluation_id}")
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Get evaluation details
        cursor.execute(
            """
            SELECT resume_path, job_title, job_description, profile_summary 
            FROM evaluations 
            WHERE id = ?
            """,
            (evaluation_id,)
        )
        eval_result = cursor.fetchone()
        
        if not eval_result:
            logging.warning(f"No evaluation found with ID: {evaluation_id}")
            return jsonify({'error': 'Evaluation not found'}), 404
        
        # Extract resume text
        resume_path = eval_result[0]
        job_description = eval_result[2]
        profile_summary = eval_result[3]
        
        if not resume_path:
            return jsonify({'error': 'No resume path found for this evaluation'}), 400
        
        resume_text = extract_text_from_file(resume_path)
        if not resume_text:
            return jsonify({'error': 'Failed to extract text from resume'}), 400
        
        # Generate questions
        logging.info(f"Generating questions for resume: {resume_path}")
        questions_data = await async_generate_questions(
            resume_text,
            job_description,
            profile_summary
        )
        
        technical_questions = questions_data.get("TechnicalQuestions", [])
        nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
        behavioral_questions = QUICK_CHECKS
        
        # Save questions to database
        try:
            # First check if there's an existing entry
            cursor.execute(
                "SELECT id FROM interview_questions WHERE evaluation_id = ?",
                (evaluation_id,)
            )
            existing = cursor.fetchone()
            
            if existing:
                # Update existing entry
                cursor.execute(
                    """
                    UPDATE interview_questions 
                    SET technical_questions = ?,
                        nontechnical_questions = ?,
                        behavioral_questions = ?
                    WHERE evaluation_id = ?
                    """,
                    (
                        json.dumps(technical_questions), 
                        json.dumps(nontechnical_questions), 
                        json.dumps(behavioral_questions), 
                        evaluation_id
                    )
                )
            else:
                # Insert new entry
                cursor.execute(
                    """
                    INSERT INTO interview_questions 
                    (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (
                        evaluation_id,
                        json.dumps(technical_questions),
                        json.dumps(nontechnical_questions),
                        json.dumps(behavioral_questions)
                    )
                )
            
            conn.commit()
            logging.info(f"Saved questions for evaluation ID: {evaluation_id}")
        except Exception as e:
            logging.error(f"Error saving questions to database: {str(e)}")
            conn.rollback()
        
        # Return the generated questions
        return jsonify({
            'technical_questions': technical_questions,
            'nontechnical_questions': nontechnical_questions,
            'behavioral_questions': behavioral_questions
        })
        
    except Exception as e:
        logging.error(f"Error generating questions: {str(e)}")
        return jsonify({'error': str(e)}), 500
        
    finally:
        if conn:
            conn.close()

@app.route('/api/generate-recruiter-handbook', methods=['POST'])
async def generate_recruiter_handbook():
    """API endpoint to generate a comprehensive recruiter handbook"""
    try:
        data = request.get_json()
        job_title = data.get('job_title', '').strip()
        job_description = data.get('job_description', '').strip()
        additional_context = data.get('additional_context', '').strip()
        oorwin_job_id = data.get('oorwin_job_id', '').strip()
        
        if not job_title:
            return jsonify({
                'success': False,
                'message': 'Job title is required'
            }), 400
            
        if not job_description:
            return jsonify({
                'success': False,
                'message': 'Job description is required'
            }), 400
        
        logging.info(f"Generating recruiter handbook for JobID: {oorwin_job_id or 'None'}...")
        
        # Create the prompt for Gemini
        handbook_prompt = f"""You are an expert recruitment specialist with deep experience in creating recruiter playbooks and handbooks for various roles across industries. Your task is to analyze a provided Job Description (JD) and generate a comprehensive "Recruiter Playbook & Handbook" in a structured, professional format that mirrors the style, structure, and content of the example playbooks you have been trained on (e.g., for roles like Head of Engineering at Fractal, Product Manager – Monogastric at Jubilant Ingrevia, and Senior Research Scientist – Analytical Chemistry at Jubilant Ingrevia).

Key guidelines for the output:

**CRITICAL ORDER - Follow this exact sequence (NO duplicates, NO extra sections):**

1. **Title**: Start with an emoji like 📖, followed by "Recruiter Playbook & Handbook: [Role Title] ([Company Name])".

2. **Mini Table of Contents**: Immediately after the title, add a compact "Mini Table of Contents" with markdown links to each main section (H2 headings). Keep it to one line per section. Format as: "- [Section Name](#section-name)"

3. **Introduction**: A brief paragraph (ONLY ONE paragraph) explaining that the handbook equips recruiters with JD analysis, screening framework, sourcing tactics, red flags, a recruiter sales pitch, and more to engage candidates effectively. Include the job link if provided in the JD; otherwise, omit it. Do NOT repeat the title or add any other content before this introduction.

4. Structure the content exactly as follows (use numbered sections, bullet points, and sub-bullets for clarity; incorporate emojis like ✨ for sales pitch, ✅ for closing, • for lists):

   **1. Primary Sourcing Parameters (Must-Have)**: Produce a compact, scannable table using GitHub-Flavored Markdown with EXACTLY these columns and order:
      | # | Skill / Experience | Recruiter Cue | Why It Matters |
      Then add 6–8 rows, where:
      - "#" is a running number starting at 1 (plain numbers only)
      - "Skill / Experience" is the must-have capability (e.g., Microfrontend, React + Next.js/Remix, Frontend Platform Development, System Design, Tech Leadership & Mentorship, Strong Communication)
      - "Recruiter Cue" lists concrete signals/cues (comma-separated; keep concise, e.g., "Module Federation, SingleSPA, NX, Turborepo")
      - "Why It Matters" explains impact/value in a short phrase (one line)
      CRITICAL: Output MUST be a valid markdown table with header separator row (|---|---|---|---|). No extra prose between the H2 and the table.

   **2. Screening Framework**: Categorize into sections like A. Background & Motivation, B. Domain Experience, etc. (up to G. Practicalities). Each section should have 1-3 bullet-pointed screening questions derived from the JD. Keep questions open-ended and probing.

   **3. Target Talent Pools**:
      - **Likely Companies**: List 4-8 relevant companies, ONE per line as separate bullets. Do NOT put multiple companies on one line.
      - **Likely Titles**: List 3-5 alternative job titles, ONE per line as separate bullets. Do NOT put multiple titles on one line.
      - **Boolean Search Samples**: Provide **EXACTLY 3 DIFFERENT Boolean search strings**, each one **STRICTLY under 200 characters**. Format each as follows:
        * **Sample 1 (Skills-focused):** `[your boolean string here]`
        * **Sample 2 (Company-focused):** `[your boolean string here]`
        * **Sample 3 (Title-focused):** `[your boolean string here]`
        CRITICAL: Each string MUST be under 200 characters. Count characters carefully. Focus on key terms only.

   **4. Red Flags to Watch**: List 4-6 bullet-pointed red flags (e.g., lack of specific experience) based on potential mismatches from the JD.

   **5. Recruiter Sales Pitch (to candidates)**: Start with ✨ **Why [Company]?** List 5-7 bullet points highlighting company strengths, role impact, growth opportunities, etc. Infer from the JD or common knowledge; make it engaging and positive. End with a closing tagline.

   **6. Recruiter Checklist (Pre-call)**: List 4-6 bullet points of key pre-call actions (e.g., confirm experience, probe specifics).

   **7. Overqualification/Overkill Risk Assessment**: ⚠️ **CRITICAL** - Analyze if candidates for this role might be overqualified or "overkill". Consider:
      - **Experience Level Mismatch**: If JD asks for 3-5 years but attracts 10+ year candidates
      - **Title Level Gap**: If role is Manager level but attracts Director/VP candidates  
      - **Compensation Concerns**: If role budget might not match senior candidates' expectations
      - **Flight Risk Indicators**: Red flags that candidate may leave quickly (overqualified, career plateau, lateral move)
      - **When to Proceed Anyway**: Scenarios where overqualified candidates are worth considering (career change, lifestyle choice, genuine interest)
      - **Screening Questions**: 2-3 specific questions to probe motivation and flight risk for overqualified candidates
      
      End with ✅ followed by: "This handbook provides recruiters with JD analysis, structured screening questions, sourcing pools, red flags, and overqualification assessment. Use as a starting document and conduct your own research before commencing the search."

**Style**: Professional, concise, actionable. Use bold for section headers. Incorporate industry-specific nuances from the JD. Assume good intent and focus on fit. Do not add unrelated content. Output in markdown format for readability.

---

**Job Description:**
{job_description}

{"**Additional Context:**" if additional_context else ""}
{additional_context if additional_context else ""}

---

Generate the complete Recruiter Playbook & Handbook now:"""

        # Generate handbook using Gemini
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(handbook_prompt)
        
        if not response or not response.text:
            raise Exception("Failed to generate handbook content from AI")
        
        handbook_content = response.text
        
        logging.info("Recruiter handbook generated successfully")
        
        # Save handbook to database
        try:
            conn = sqlite3.connect('combined_db.db')
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT INTO recruiter_handbooks (
                    oorwin_job_id, job_title, job_description, 
                    additional_context, markdown_content, timestamp
                ) VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                oorwin_job_id if oorwin_job_id else None,
                job_title,
                job_description,
                additional_context,
                handbook_content,
                datetime.now()
            ))
            
            conn.commit()
            handbook_id = cursor.lastrowid
            conn.close()
            
            logging.info(f"Handbook saved to database with ID: {handbook_id}")
            
            # Return success with handbook_id
            return jsonify({
                'success': True,
                'markdown_content': handbook_content,
                'handbook_id': handbook_id
            })
        except Exception as e:
            logging.error(f"Error saving handbook to database: {str(e)}")
            # Return without handbook_id if save fails
            return jsonify({
                'success': True,
                'markdown_content': handbook_content
            })
        
    except Exception as e:
        logging.error(f"Error generating recruiter handbook: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/get-job-ids', methods=['GET'])
def get_job_ids():
    """API endpoint to get all unique JobIDs for auto-suggest"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Get unique JobIDs from evaluations
        cursor.execute('''
            SELECT DISTINCT oorwin_job_id 
            FROM evaluations 
            WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ''
            ORDER BY oorwin_job_id
        ''')
        eval_job_ids = [row[0] for row in cursor.fetchall()]
        
        # Get unique JobIDs from handbooks
        cursor.execute('''
            SELECT DISTINCT oorwin_job_id 
            FROM recruiter_handbooks 
            WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ''
            ORDER BY oorwin_job_id
        ''')
        handbook_job_ids = [row[0] for row in cursor.fetchall()]
        
        # Combine and remove duplicates
        all_job_ids = list(set(eval_job_ids + handbook_job_ids))
        all_job_ids.sort()
        
        conn.close()
        
        return jsonify({
            'success': True,
            'job_ids': all_job_ids
        })
        
    except Exception as e:
        logging.error(f"Error fetching JobIDs: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/get-job-data/<job_id>', methods=['GET'])
def get_job_data(job_id):
    """API endpoint to get job description for auto-fill based on JobID"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # First try to get from handbooks (they always have full JD)
        cursor.execute('''
            SELECT job_title, job_description, additional_context 
            FROM recruiter_handbooks 
            WHERE oorwin_job_id = ? 
            ORDER BY timestamp DESC 
            LIMIT 1
        ''', (job_id,))
        
        result = cursor.fetchone()
        
        if result and result[1]:  # Check if job_description is not empty
            conn.close()
            return jsonify({
                'success': True,
                'job_title': result[0],
                'job_description': result[1],
                'additional_context': result[2] if result[2] else '',
                'source': 'handbook'
            })
        
        # If not found in handbooks or JD is empty, try evaluations
        cursor.execute('''
            SELECT job_title, job_description 
            FROM evaluations 
            WHERE oorwin_job_id = ? AND job_description IS NOT NULL AND job_description != ''
            ORDER BY timestamp DESC 
            LIMIT 1
        ''', (job_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        if result:
            return jsonify({
                'success': True,
                'job_title': result[0],
                'job_description': result[1],
                'source': 'evaluation'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'No data found for this JobID'
            }), 404
        
    except Exception as e:
        logging.error(f"Error fetching job data for {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/get-handbooks', methods=['GET'])
def get_handbooks():
    """API endpoint to get all recruiter handbooks"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, oorwin_job_id, job_title, job_description, 
                   additional_context, markdown_content, timestamp
            FROM recruiter_handbooks
            ORDER BY timestamp DESC
        ''')
        
        rows = cursor.fetchall()
        conn.close()
        
        handbooks = []
        for row in rows:
            handbooks.append({
                'id': row[0],
                'oorwin_job_id': row[1],
                'job_title': row[2],
                'job_description': row[3],
                'additional_context': row[4],
                'markdown_content': row[5],
                'timestamp': row[6]
            })
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

# ===== DASHBOARD & ANALYTICS ROUTES =====

@app.route('/dashboard')
def dashboard():
    """Analytics dashboard showing usage metrics"""
    return render_template('dashboard.html')

@app.route('/api/analytics/overview', methods=['GET'])
def get_analytics_overview():
    """API endpoint for dashboard overview metrics"""
    try:
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Get total counts
        cursor.execute('SELECT COUNT(*) FROM evaluations')
        total_evaluations = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM recruiter_handbooks')
        total_handbooks = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(DISTINCT oorwin_job_id) FROM evaluations WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ""')
        unique_jobs_evals = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(DISTINCT oorwin_job_id) FROM recruiter_handbooks WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ""')
        unique_jobs_handbooks = cursor.fetchone()[0]
        
        total_jobs = max(unique_jobs_evals, unique_jobs_handbooks)
        
        # Get average match score
        cursor.execute('SELECT AVG(match_percentage) FROM evaluations WHERE match_percentage IS NOT NULL')
        avg_match_score = cursor.fetchone()[0] or 0
        
        # Get conversion rate (jobs with handbook that also have evaluations)
        cursor.execute('''
            SELECT COUNT(DISTINCT h.oorwin_job_id) 
            FROM recruiter_handbooks h
            INNER JOIN evaluations e ON h.oorwin_job_id = e.oorwin_job_id
            WHERE h.oorwin_job_id IS NOT NULL AND h.oorwin_job_id != ""
        ''')
        jobs_with_both = cursor.fetchone()[0]
        conversion_rate = (jobs_with_both / total_handbooks * 100) if total_handbooks > 0 else 0
        
        # Get average evaluations per job
        cursor.execute('''
            SELECT AVG(eval_count) FROM (
                SELECT oorwin_job_id, COUNT(*) as eval_count 
                FROM evaluations 
                WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ""
                GROUP BY oorwin_job_id
            )
        ''')
        avg_evals_per_job = cursor.fetchone()[0] or 0
        
        # Get active jobs (activity in last 7 days)
        cursor.execute('''
            SELECT COUNT(DISTINCT oorwin_job_id) 
            FROM evaluations 
            WHERE oorwin_job_id IS NOT NULL 
            AND oorwin_job_id != ""
            AND datetime(timestamp) >= datetime('now', '-7 days')
        ''')
        active_jobs_7d = cursor.fetchone()[0]
        
        # Get average eval time (placeholder - would need tracking)
        avg_eval_time = 22  # seconds - placeholder
        
        conn.close()
        
        return jsonify({
            'success': True,
            'metrics': {
                'total_evaluations': total_evaluations,
                'total_handbooks': total_handbooks,
                'total_jobs': total_jobs,
                'active_jobs': active_jobs_7d,
                'avg_match_score': round(avg_match_score, 1),
                'conversion_rate': round(conversion_rate, 1),
                'avg_evals_per_job': round(avg_evals_per_job, 1),
                'avg_eval_time': avg_eval_time
            }
        })
        
    except Exception as e:
        logging.error(f"Error fetching analytics overview: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/timeline', methods=['GET'])
def get_analytics_timeline():
    """API endpoint for activity timeline chart data"""
    try:
        days = int(request.args.get('days', 30))
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        # Get daily evaluation counts
        cursor.execute(f'''
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM evaluations
            WHERE datetime(timestamp) >= datetime('now', '-{days} days')
            GROUP BY DATE(timestamp)
            ORDER BY date
        ''')
        eval_data = cursor.fetchall()
        
        # Get daily handbook counts
        cursor.execute(f'''
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM recruiter_handbooks
            WHERE datetime(timestamp) >= datetime('now', '-{days} days')
            GROUP BY DATE(timestamp)
            ORDER BY date
        ''')
        handbook_data = cursor.fetchall()
        
        conn.close()
        
        # Convert to dict for easy merging
        eval_dict = {row[0]: row[1] for row in eval_data}
        handbook_dict = {row[0]: row[1] for row in handbook_data}
        
        # Merge and create timeline
        from datetime import datetime, timedelta
        timeline = []
        start_date = datetime.now() - timedelta(days=days)
        
        for i in range(days + 1):
            current_date = (start_date + timedelta(days=i)).strftime('%Y-%m-%d')
            timeline.append({
                'date': current_date,
                'evaluations': eval_dict.get(current_date, 0),
                'handbooks': handbook_dict.get(current_date, 0)
            })
        
        return jsonify({
            'success': True,
            'timeline': timeline
        })
        
    except Exception as e:
        logging.error(f"Error fetching timeline data: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/analytics/top-jobs', methods=['GET'])
def get_top_jobs():
    """API endpoint for top jobs by activity"""
    try:
        limit = int(request.args.get('limit', 10))
        conn = sqlite3.connect('combined_db.db')
        cursor = conn.cursor()
        
        cursor.execute(f'''
            SELECT 
                e.oorwin_job_id,
                e.job_title,
                COUNT(*) as eval_count,
                AVG(e.match_percentage) as avg_score,
                MAX(e.timestamp) as last_active
            FROM evaluations e
            WHERE e.oorwin_job_id IS NOT NULL AND e.oorwin_job_id != ""
            GROUP BY e.oorwin_job_id, e.job_title
            ORDER BY eval_count DESC
            LIMIT {limit}
        ''')
        
        rows = cursor.fetchall()
        conn.close()
        
        jobs = []
        for row in rows:
            jobs.append({
                'job_id': row[0],
                'job_title': row[1],
                'eval_count': row[2],
                'avg_score': round(row[3], 1) if row[3] else 0,
                'last_active': row[4]
            })
        
        return jsonify({
            'success': True,
            'jobs': jobs
        })
        
    except Exception as e:
        logging.error(f"Error fetching top jobs: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/handbook/<int:handbook_id>', methods=['GET'])
def get_single_handbook(handbook_id):
    """API endpoint to get a single handbook by ID"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, job_title, oorwin_job_id, job_description, timestamp, markdown_content
            FROM recruiter_handbooks
            WHERE id = ?
        ''', (handbook_id,))
        
        row = cursor.fetchone()
        conn.close()
        
        if not row:
            return jsonify({
                'success': False,
                'message': 'Handbook not found'
            }), 404
        
        handbook = {
            'id': row[0],
            'job_title': row[1],
            'oorwin_job_id': row[2],
            'job_description': row[3],
            'timestamp': row[4],
            'markdown_content': row[5]
        }
        
        return jsonify({
            'success': True,
            'handbook': handbook
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbook {handbook_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/handbooks-only', methods=['GET'])
def get_handbooks_only():
    """API endpoint for handbooks-only history"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, job_title, oorwin_job_id, timestamp, markdown_content
            FROM recruiter_handbooks
            ORDER BY timestamp DESC
        ''')
        
        handbooks = []
        for row in cursor.fetchall():
            handbooks.append({
                'id': row[0],
                'job_title': row[1],
                'oorwin_job_id': row[2],
                'timestamp': row[3],
                'markdown_content': row[4][:500] + '...' if row[4] and len(row[4]) > 500 else row[4]  # Truncate for preview
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks-only history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/evaluations-only', methods=['GET'])
def get_evaluations_only():
    """API endpoint for evaluations-only history"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, filename, job_title, oorwin_job_id, match_percentage, timestamp
            FROM evaluations
            ORDER BY timestamp DESC
        ''')
        
        evaluations = []
        for row in cursor.fetchall():
            evaluations.append({
                'id': row[0],
                'filename': row[1],
                'job_title': row[2],
                'oorwin_job_id': row[3] if row[3] else 'N/A',
                'match_percentage': row[4],
                'timestamp': row[5]
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'evaluations': evaluations
        })
        
    except Exception as e:
        logging.error(f"Error fetching evaluations-only history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/job-centric-history', methods=['GET'])
def get_job_centric_history():
    """API endpoint for job-centric grouped history"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get all unique job IDs from both tables, ordered by most recent activity
        cursor.execute('''
            SELECT oorwin_job_id, MAX(timestamp) as last_activity
            FROM (
                SELECT oorwin_job_id, timestamp FROM evaluations WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ""
                UNION ALL
                SELECT oorwin_job_id, timestamp FROM recruiter_handbooks WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ""
            )
            GROUP BY oorwin_job_id
            ORDER BY last_activity DESC
        ''')
        
        job_ids = [row[0] for row in cursor.fetchall()]
        
        result = []
        for job_id in job_ids:
            # Get job title (prefer from handbooks, fallback to evaluations)
            cursor.execute('''
                SELECT job_title FROM recruiter_handbooks 
                WHERE oorwin_job_id = ? AND job_title IS NOT NULL AND job_title != ""
                ORDER BY timestamp DESC LIMIT 1
            ''', (job_id,))
            
            job_title_row = cursor.fetchone()
            job_title = job_title_row[0] if job_title_row else None
            
            if not job_title:
                cursor.execute('''
                    SELECT job_title FROM evaluations 
                    WHERE oorwin_job_id = ? AND job_title IS NOT NULL AND job_title != ""
                    ORDER BY timestamp DESC LIMIT 1
                ''', (job_id,))
                job_title_row = cursor.fetchone()
                job_title = job_title_row[0] if job_title_row else "N/A"
            
            # Count handbooks
            cursor.execute('''
                SELECT COUNT(*) FROM recruiter_handbooks WHERE oorwin_job_id = ?
            ''', (job_id,))
            handbooks_count = cursor.fetchone()[0]
            
            # Get evaluated resumes (filenames)
            cursor.execute('''
                SELECT filename, timestamp, match_percentage FROM evaluations 
                WHERE oorwin_job_id = ?
                ORDER BY timestamp DESC
            ''', (job_id,))
            evaluations = cursor.fetchall()
            evaluations_count = len(evaluations)
            
            resume_list = []
            for eval_row in evaluations:
                resume_list.append({
                    'filename': eval_row[0],
                    'timestamp': eval_row[1],
                    'match_percentage': eval_row[2]
                })
            
            # Get first created date (earliest timestamp from both tables)
            cursor.execute('''
                SELECT MIN(timestamp) FROM (
                    SELECT timestamp FROM evaluations WHERE oorwin_job_id = ?
                    UNION ALL
                    SELECT timestamp FROM recruiter_handbooks WHERE oorwin_job_id = ?
                )
            ''', (job_id, job_id))
            first_created = cursor.fetchone()[0]
            
            # Get last activity date (latest timestamp from both tables)
            cursor.execute('''
                SELECT MAX(timestamp) FROM (
                    SELECT timestamp FROM evaluations WHERE oorwin_job_id = ?
                    UNION ALL
                    SELECT timestamp FROM recruiter_handbooks WHERE oorwin_job_id = ?
                )
            ''', (job_id, job_id))
            last_activity = cursor.fetchone()[0]
            
            result.append({
                'job_id': job_id,
                'job_title': job_title,
                'handbooks_count': handbooks_count,
                'evaluations_count': evaluations_count,
                'resume_list': resume_list,
                'first_created': first_created,
                'last_activity': last_activity
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'jobs': result
        })
        
    except Exception as e:
        logging.error(f"Error fetching job-centric history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/handbooks-by-job/<job_id>', methods=['GET'])
def get_handbooks_by_job(job_id):
    """API endpoint to get all handbooks for a specific job ID"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, job_title, job_description, additional_context, markdown_content, timestamp
            FROM recruiter_handbooks
            WHERE oorwin_job_id = ?
            ORDER BY timestamp DESC
        ''', (job_id,))
        
        rows = cursor.fetchall()
        conn.close()
        
        handbooks = []
        for row in rows:
            handbooks.append({
                'id': row[0],
                'job_title': row[1],
                'job_description': row[2],
                'additional_context': row[3],
                'markdown_content': row[4],
                'timestamp': row[5]
            })
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks for job {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/evaluation-full/<int:eval_id>', methods=['GET'])
def get_evaluation_full(eval_id):
    """API endpoint to get full evaluation data for viewing"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Fetch evaluation data
        cursor.execute('''
            SELECT 
                e.id, e.filename, e.job_title, e.job_description,
                e.match_percentage, e.match_factors, e.profile_summary,
                e.missing_keywords, e.job_stability, e.career_progression,
                e.oorwin_job_id, e.timestamp, e.candidate_fit_analysis, e.over_under_qualification
            FROM evaluations e
            WHERE e.id = ?
        ''', (eval_id,))
        
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return jsonify({
                'success': False,
                'message': 'Evaluation not found'
            }), 404
        
        # Fetch interview questions
        cursor.execute('''
            SELECT technical_questions, nontechnical_questions, behavioral_questions
            FROM interview_questions
            WHERE evaluation_id = ?
        ''', (eval_id,))
        
        questions_row = cursor.fetchone()
        conn.close()
        
        # Parse JSON fields
        import json
        evaluation = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'job_description': row[3],
            'match_percentage': row[4],
            'match_percentage_str': str(int(row[4])) + '%' if row[4] else '0%',
            'match_factors': json.loads(row[5]) if row[5] else {},
            'profile_summary': row[6],
            'missing_keywords': json.loads(row[7]) if row[7] else [],
            'job_stability': json.loads(row[8]) if row[8] else {},
            'career_progression': json.loads(row[9]) if row[9] else {},
            'oorwin_job_id': row[10],
            'timestamp': row[11],
            # Parse new fields from database (or use empty if not present)
            'candidate_fit_analysis': json.loads(row[12]) if (len(row) > 12 and row[12]) else {},
            'over_under_qualification': row[13] if (len(row) > 13 and row[13]) else ''
        }
        
        if questions_row:
            evaluation['technical_questions'] = json.loads(questions_row[0]) if questions_row[0] else []
            evaluation['nontechnical_questions'] = json.loads(questions_row[1]) if questions_row[1] else []
            evaluation['behavioral_questions'] = json.loads(questions_row[2]) if questions_row[2] else []
        else:
            evaluation['technical_questions'] = []
            evaluation['nontechnical_questions'] = []
            evaluation['behavioral_questions'] = []
        
        return jsonify({
            'success': True,
            'evaluation': evaluation
        })
        
    except Exception as e:
        logging.error(f"Error fetching full evaluation {eval_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/evaluations-by-job/<job_id>', methods=['GET'])
def get_evaluations_by_job(job_id):
    """API endpoint to get all evaluations for a specific job ID"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, filename, match_percentage, timestamp
            FROM evaluations
            WHERE oorwin_job_id = ?
            ORDER BY timestamp DESC
        ''', (job_id,))
        
        rows = cursor.fetchall()
        conn.close()
        
        evaluations = []
        for row in rows:
            evaluations.append({
                'id': row[0],
                'filename': row[1],
                'match_percentage': row[2],
                'timestamp': row[3]
            })
        
        return jsonify({
            'success': True,
            'evaluations': evaluations
        })
        
    except Exception as e:
        logging.error(f"Error fetching evaluations for job {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/download-handbook-pdf', methods=['POST'])
async def download_handbook_pdf():
    """API endpoint to download recruiter handbook as PDF"""
    try:
        data = request.get_json()
        markdown_content = data.get('markdown_content', '').strip()
        job_title = (data.get('job_title') or '').strip()
        oorwin_job_id = (data.get('oorwin_job_id') or '').strip()
        
        if not markdown_content:
            return jsonify({
                'success': False,
                'message': 'No content to download'
            }), 400
        
        logging.info("Generating PDF from handbook content...")
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=50,  # Reduced margins for better table fit
            leftMargin=50,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#2c3e50',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#2c3e50',
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Blue style for main section titles
        heading1_blue_style = ParagraphStyle(
            'CustomHeading1Blue',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#3498db',  # Blue color
            spaceAfter=12,
            spaceBefore=12
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#34495e',
            spaceAfter=10,
            spaceBefore=10
        )
        
        # Blue style for main section titles (H2)
        heading2_blue_style = ParagraphStyle(
            'CustomHeading2Blue',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#3498db',  # Blue color
            spaceAfter=10,
            spaceBefore=10
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['BodyText'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=6
        )
        
        # Helper function to check if heading is a main section title (should be blue)
        def is_main_section_title(text):
            """Check if heading text matches one of the 7 main section titles (EXCLUDE Introduction)"""
            # Explicitly exclude "Introduction"
            text_clean = re.sub(r'[^\x00-\x7F]+', '', text).strip()  # Remove emojis first
            if re.match(r'^Introduction\s*:?\s*$', text_clean, re.IGNORECASE):
                return False
            
            main_titles = [
                r'^\s*\d+\.?\s*Primary\s+Sourcing\s+Parameters\s*\(Must-Have\)\s*:?',
                r'^\s*\d+\.?\s*Screening\s+Framework\s*:?',
                r'^\s*\d+\.?\s*Target\s+Talent\s+Pools\s*:?',
                r'^\s*\d+\.?\s*Red\s+Flags\s+to\s+Watch\s*:?',
                r'^\s*\d+\.?\s*Recruiter\s+Sales\s+Pitch\s*\(to\s+candidates\)\s*:?',
                r'^\s*\d+\.?\s*Recruiter\s+Checklist\s*\(Pre-call\)\s*:?',
                r'^\s*\d+\.?\s*Overqualification/Overkill\s+Risk\s+Assessment\s*:?',
                # Also match without numbers (in case LLM doesn't include them)
                r'^\s*Primary\s+Sourcing\s+Parameters\s*\(Must-Have\)\s*:?',
                r'^\s*Screening\s+Framework\s*:?',
                r'^\s*Target\s+Talent\s+Pools\s*:?',
                r'^\s*Red\s+Flags\s+to\s+Watch\s*:?',
                r'^\s*Recruiter\s+Sales\s+Pitch\s*\(to\s+candidates\)\s*:?',
                r'^\s*Recruiter\s+Checklist\s*\(Pre-call\)\s*:?',
                r'^\s*Overqualification/Overkill\s+Risk\s+Assessment\s*:?'
            ]
            return any(re.match(pattern, text_clean, re.IGNORECASE) for pattern in main_titles)
        
        # Parse markdown content and convert to PDF elements
        lines = markdown_content.split('\n')
        i = 0
        in_table = False
        table_rows = []
        seen_intro = False  # Track if we've seen Introduction section
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines (but add small spacing)
            if not line:
                if not in_table:
                    elements.append(Spacer(1, 0.15*inch))
                i += 1
                continue
            
            # Skip duplicate "Introduction:" lines (blue colored duplicates that appear after title)
            if re.match(r'^Introduction:?\s*$', line, re.IGNORECASE):
                if seen_intro:
                    # Skip this duplicate intro line and any following empty lines
                    i += 1
                    while i < len(lines) and not lines[i].strip():
                        i += 1
                    continue
                seen_intro = True
            
            # Skip TOC markdown links (they don't work in PDF anyway)
            if re.match(r'^-\s*\[.*\]\(#.*\)', line):
                i += 1
                continue
            
            # Handle markdown tables
            if '|' in line and line.count('|') >= 2:
                # Check if it's a table header separator (contains dashes)
                if re.match(r'^\|[\s\-:]+\|', line):
                    i += 1
                    continue
                
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove first/last empty
                if cells:
                    table_rows.append(cells)
                    in_table = True
                    i += 1
                    continue
            else:
                # If we were building a table, finish it
                if in_table and table_rows:
                    # Create PDF table
                    table_data = table_rows
                    
                    # Check if this is "Primary Sourcing Parameters" table (first column should be blue)
                    is_sourcing_table = len(table_data) > 0 and len(table_data[0]) > 0 and (
                        'Skill' in str(table_data[0]) or 'Experience' in str(table_data[0]) or 
                        'Recruiter Cue' in str(table_data[0])
                    )
                    
                    # Calculate available width (A4 width - margins in points)
                    available_width = A4[0] - 100  # 50pt margin on each side (72pt = 1 inch)
                    
                    # Calculate column widths based on number of columns
                    num_cols = len(table_data[0]) if table_data else 4
                    if is_sourcing_table and num_cols >= 4:
                        # For Primary Sourcing Parameters: narrow first col, distribute rest
                        col_widths = [35] + [(available_width - 35) / (num_cols - 1)] * (num_cols - 1)
                    else:
                        # Equal width for all columns
                        col_widths = [available_width / num_cols] * num_cols
                    
                    # Convert table cells to Paragraph objects for proper text wrapping
                    table_style_small = ParagraphStyle(
                        'TableCell',
                        parent=styles['Normal'],
                        fontSize=9,
                        leading=11,  # Line height
                        spaceAfter=0,
                        spaceBefore=0,
                    )
                    
                    # Process table data: convert strings to Paragraphs for wrapping
                    table_style_header = ParagraphStyle(
                        'TableHeader',
                        parent=styles['Normal'],
                        fontSize=10,
                        leading=12,
                        spaceAfter=0,
                        spaceBefore=0,
                        fontName='Helvetica-Bold',
                    )
                    
                    processed_table_data = []
                    for row_idx, row in enumerate(table_data):
                        processed_row = []
                        for col_idx, cell_text in enumerate(row):
                            # Escape HTML and convert to Paragraph for wrapping
                            cell_text_clean = str(cell_text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            # Convert markdown bold/italic
                            cell_text_clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', cell_text_clean)
                            cell_text_clean = re.sub(r'\*(.+?)\*', r'<i>\1</i>', cell_text_clean)
                            # Use header style for first row, body style for others
                            cell_style = table_style_header if row_idx == 0 else table_style_small
                            # Create Paragraph for proper wrapping
                            processed_row.append(Paragraph(cell_text_clean, cell_style))
                        processed_table_data.append(processed_row)
                    
                    pdf_table = Table(processed_table_data, colWidths=col_widths)
                    
                    # Build table style
                    table_style = [
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),  # Dark header
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),  # Header font size
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                        ('TOPPADDING', (0, 0), (-1, 0), 6),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#34495e')),  # Thinner grid
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Top align for multi-line cells
                        ('LEFTPADDING', (0, 0), (-1, -1), 5),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 4),  # Padding for body cells
                        ('TOPPADDING', (0, 1), (-1, -1), 4),
                    ]
                    
                    # If it's the Primary Sourcing Parameters table, style first column (numbers) in blue
                    if is_sourcing_table and len(table_data[0]) >= 4:
                        table_style.extend([
                            ('TEXTCOLOR', (0, 1), (0, -1), colors.HexColor('#3498db')),  # Blue for first column
                            ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),  # Bold first column
                            ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # Center first column
                        ])
                    
                    pdf_table.setStyle(TableStyle(table_style))
                    elements.append(Spacer(1, 0.1*inch))
                    elements.append(pdf_table)
                    elements.append(Spacer(1, 0.2*inch))
                    table_rows = []
                    in_table = False
            
            # Escape special characters for ReportLab
            line = unescape(line)
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Handle headers
            if line.startswith('# '):
                text = line[2:].strip()
                # Remove emojis from headers
                text = re.sub(r'[^\x00-\x7F]+', '', text).strip()
                elements.append(Spacer(1, 0.2*inch))
                elements.append(Paragraph(text, title_style))
            elif line.startswith('## '):
                text = line[3:].strip()
                text_clean = re.sub(r'[^\x00-\x7F]+', '', text).strip()
                elements.append(Spacer(1, 0.15*inch))
                # Use blue style if it's a main section title
                style_to_use = heading1_blue_style if is_main_section_title(text_clean) else heading1_style
                elements.append(Paragraph(text_clean, style_to_use))
            elif line.startswith('### '):
                text = line[4:].strip()
                text_clean = re.sub(r'[^\x00-\x7F]+', '', text).strip()
                elements.append(Spacer(1, 0.1*inch))
                # Use blue style if it's a main section title
                style_to_use = heading2_blue_style if is_main_section_title(text_clean) else heading2_style
                elements.append(Paragraph(text_clean, style_to_use))
            elif line.startswith('#### '):
                text = line[5:].strip()
                text_clean = re.sub(r'[^\x00-\x7F]+', '', text).strip()
                elements.append(Spacer(1, 0.08*inch))
                # Use blue style if it's a main section title
                style_to_use = heading2_blue_style if is_main_section_title(text_clean) else heading2_style
                elements.append(Paragraph(text_clean, style_to_use))
            # Handle bullet points
            elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                text = line[2:].strip() if line.startswith('- ') or line.startswith('* ') else line[2:].strip()
                # Convert markdown formatting in bullets
                text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                elements.append(Paragraph('• ' + text, bullet_style))
            elif line.startswith('o '):
                text = line[2:].strip()
                text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                elements.append(Paragraph('○ ' + text, bullet_style))
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                elements.append(Paragraph(text, bullet_style))
            # Handle bold text (standalone)
            elif line.startswith('**') and line.endswith('**'):
                text = '<b>' + line[2:-2].strip() + '</b>'
                elements.append(Paragraph(text, body_style))
            # Handle horizontal rules
            elif line.startswith('---') or line.startswith('___'):
                elements.append(Spacer(1, 0.3*inch))
            # Regular text
            else:
                # Convert markdown bold and italic
                text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
                text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                # Remove emojis from regular text (optional - comment out if you want to keep them)
                # text = re.sub(r'[^\x00-\x7F]+', '', text)
                if text.strip():
                    elements.append(Paragraph(text, body_style))
            
            i += 1
        
        # Handle any remaining table
        if in_table and table_rows:
            table_data = table_rows
            
            # Check if this is "Primary Sourcing Parameters" table
            is_sourcing_table = len(table_data) > 0 and len(table_data[0]) > 0 and (
                'Skill' in str(table_data[0]) or 'Experience' in str(table_data[0]) or 
                'Recruiter Cue' in str(table_data[0])
            )
            
            # Calculate available width (A4 width - margins in points)
            available_width = A4[0] - 100  # 50pt margin on each side
            
            # Calculate column widths
            num_cols = len(table_data[0]) if table_data else 4
            if is_sourcing_table and num_cols >= 4:
                col_widths = [35] + [(available_width - 35) / (num_cols - 1)] * (num_cols - 1)
            else:
                col_widths = [available_width / num_cols] * num_cols
            
            # Convert table cells to Paragraph objects for proper text wrapping
            table_style_small = ParagraphStyle(
                'TableCell',
                parent=styles['Normal'],
                fontSize=9,
                leading=11,  # Line height
                spaceAfter=0,
                spaceBefore=0,
            )
            
            # Process table data: convert strings to Paragraphs for wrapping
            table_style_header = ParagraphStyle(
                'TableHeader',
                parent=styles['Normal'],
                fontSize=10,
                leading=12,
                spaceAfter=0,
                spaceBefore=0,
                fontName='Helvetica-Bold',
            )
            
            processed_table_data = []
            for row_idx, row in enumerate(table_data):
                processed_row = []
                for col_idx, cell_text in enumerate(row):
                    # Escape HTML and convert to Paragraph for wrapping
                    cell_text_clean = str(cell_text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    # Convert markdown bold/italic
                    cell_text_clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', cell_text_clean)
                    cell_text_clean = re.sub(r'\*(.+?)\*', r'<i>\1</i>', cell_text_clean)
                    # Use header style for first row, body style for others
                    cell_style = table_style_header if row_idx == 0 else table_style_small
                    # Create Paragraph for proper wrapping
                    processed_row.append(Paragraph(cell_text_clean, cell_style))
                processed_table_data.append(processed_row)
            
            pdf_table = Table(processed_table_data, colWidths=col_widths)
            
            table_style = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('TOPPADDING', (0, 0), (-1, 0), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#34495e')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ('TOPPADDING', (0, 1), (-1, -1), 4),
            ]
            
            if is_sourcing_table and len(table_data[0]) >= 4:
                table_style.extend([
                    ('TEXTCOLOR', (0, 1), (0, -1), colors.HexColor('#3498db')),
                    ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
                    ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                ])
            
            pdf_table.setStyle(TableStyle(table_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(pdf_table)
            elements.append(Spacer(1, 0.2*inch))
        
        # Build PDF and set metadata title
        def _set_meta(canvas, _):
            title_parts = [p for p in ["Recruiter Handbook", job_title, oorwin_job_id] if p]
            canvas.setTitle(" - ".join(title_parts))
            canvas.setAuthor("PeopleLogic PeopleBot")

        doc.build(elements, onFirstPage=_set_meta, onLaterPages=_set_meta)
        
        # Get PDF data
        pdf_data = buffer.getvalue()
        buffer.close()
        
        logging.info("PDF generated successfully")
        
        # Return PDF as response
        response = make_response(pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        
        # Format filename as RH_JobTitle_JobID.pdf
        safe_title = re.sub(r'[^A-Za-z0-9 _\-]+', '', job_title) or 'Handbook'
        safe_job = re.sub(r'[^A-Za-z0-9 _\-]+', '', oorwin_job_id) if oorwin_job_id else ''
        
        # Build filename: RH_JobTitle_JobID.pdf
        filename_parts = ["RH", safe_title]
        if safe_job:
            filename_parts.append(safe_job)
        
        filename = "_".join(filename_parts) + '.pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500


def get_default_interview_questions(job_title):
    """Generate default interview questions based on job title"""
    # Default technical questions based on common job titles
    technical_questions = {
        "software": [
            "Describe your experience with different programming languages and frameworks.",
            "How do you approach debugging a complex issue in your code?",
            "Explain your understanding of object-oriented programming principles.",
            "How do you ensure code quality and maintainability?",
            "Describe a challenging technical problem you solved recently."
        ],
        "data": [
            "Explain the difference between supervised and unsupervised learning.",
            "How do you handle missing or inconsistent data in your analysis?",
            "Describe your experience with SQL and database optimization.",
            "What tools and libraries do you use for data visualization?",
            "How do you validate the results of your data analysis?"
        ],
        "manager": [
            "How do you approach resource allocation in a project?",
            "Describe your experience with agile methodologies.",
            "How do you handle conflicts within your team?",
            "What metrics do you use to measure project success?",
            "How do you ensure your team meets deadlines and quality standards?"
        ],
        "analyst": [
            "Describe your approach to gathering requirements from stakeholders.",
            "How do you prioritize features or improvements?",
            "What tools do you use for data analysis and reporting?",
            "How do you communicate complex findings to non-technical stakeholders?",
            "Describe a situation where your analysis led to a significant business decision."
        ],
        "designer": [
            "How do you approach the design process for a new project?",
            "Describe your experience with different design tools and software.",
            "How do you incorporate user feedback into your designs?",
            "How do you balance aesthetics with functionality?",
            "Describe a design challenge you faced and how you overcame it."
        ]
    }
    
    # Default non-technical questions
    nontechnical_questions = [
        "How do you prioritize your work when dealing with multiple deadlines?",
        "Describe a situation where you had to collaborate with a difficult team member.",
        "How do you stay updated with the latest trends and developments in your field?",
        "Describe your ideal work environment and company culture.",
        "How do you handle feedback and criticism?"
    ]
    
    # Determine which set of technical questions to use based on job title
    job_title_lower = job_title.lower()
    selected_technical_questions = []
    
    if any(keyword in job_title_lower for keyword in ["developer", "engineer", "programmer", "software", "code", "web"]):
        selected_technical_questions = technical_questions["software"]
    elif any(keyword in job_title_lower for keyword in ["data", "analytics", "scientist", "ml", "ai"]):
        selected_technical_questions = technical_questions["data"]
    elif any(keyword in job_title_lower for keyword in ["manager", "director", "lead", "head"]):
        selected_technical_questions = technical_questions["manager"]
    elif any(keyword in job_title_lower for keyword in ["analyst", "business", "product"]):
        selected_technical_questions = technical_questions["analyst"]
    elif any(keyword in job_title_lower for keyword in ["designer", "ux", "ui", "graphic"]):
        selected_technical_questions = technical_questions["designer"]
    else:
        # If no match, use a mix of questions
        selected_technical_questions = [
            technical_questions["software"][0],
            technical_questions["analyst"][0],
            technical_questions["manager"][0],
            "Describe your technical skills that are most relevant to this position.",
            "What technical challenges are you looking forward to tackling in this role?"
        ]
    
    return selected_technical_questions, nontechnical_questions

# Batch evaluate multiple resumes against the same JD
@app.route('/evaluate-batch', methods=['POST'])
def evaluate_batch():
    """Evaluate multiple resumes against the same JD and return a comparison ranking."""
    try:
        if 'resumes' not in request.files:
            return jsonify({'success': False, 'error': 'No resumes provided'}), 400

        files = request.files.getlist('resumes')
        if not files:
            return jsonify({'success': False, 'error': 'No files received'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')
        if not job_title or not job_description:
            return jsonify({'success': False, 'error': 'Missing job title or description'}), 400

        results = []
        for f in files:
            if f.filename == '' or not allowed_file(f.filename):
                continue
            filename = secure_filename(f.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            f.save(file_path)

            resume_text = extract_text_from_file(file_path)
            if not resume_text:
                continue

            formatted_prompt = input_prompt_template.format(resume_text=resume_text, job_description=job_description)
            main_response = asyncio.run(async_gemini_generate(formatted_prompt))
            if not main_response:
                continue

            match_percentage_str = main_response.get('JD Match', '0%')
            match_percentage = int(str(match_percentage_str).strip('%') or 0)
            # Derive strengths from Match Factors if available (dict of dimension->score/comment)
            top_strengths = []
            mf = main_response.get('Match Factors', {})
            if isinstance(mf, dict) and mf:
                # If numeric scores, sort desc; else take top keys
                try:
                    sorted_items = sorted(mf.items(), key=lambda kv: float(str(kv[1]).split('%')[0]) if isinstance(kv[1], str) and '%' in kv[1] else float(kv[1]), reverse=True)
                except Exception:
                    sorted_items = list(mf.items())
                for k, v in sorted_items[:5]:
                    top_strengths.append(f"{k}")
            # Fallback: use extracted keywords from profile summary heuristics
            if not top_strengths:
                ps = main_response.get('Profile Summary', '') or ''
                words = [w.strip('.,;:()').title() for w in ps.split() if len(w) > 3]
                uniq = []
                for w in words:
                    if w not in uniq:
                        uniq.append(w)
                top_strengths = uniq[:5]

            key_gaps = list(main_response.get('MissingKeywords', [])) if isinstance(main_response.get('MissingKeywords', []), list) else []

            results.append({
                'filename': filename,
                'match_percentage': match_percentage,
                'top_strengths': top_strengths,
                'key_gaps': key_gaps
            })

        results.sort(key=lambda x: x['match_percentage'], reverse=True)
        if not results:
            return jsonify({'success': False, 'error': 'Failed to evaluate uploaded resumes'}), 500

        # Build a recruiter-style markdown comparison report (compact)
        def eval_mark(mark_score):
            if mark_score >= 75:
                return '✅'
            if mark_score >= 55:
                return '⚠️'
            return '❌'

        # JD Summary placeholder (kept short)
        md_lines = []
        md_lines.append('# 🧭 JD Summary')
        md_lines.append('(Concise summary of the role. Auto-generated placeholders — edit as needed.)')
        md_lines.append('')
        md_lines.append('| JD Pillar | Key Expectations |')
        md_lines.append('|------------|------------------|')
        md_lines.append('| Role Objective | Define and deliver measurable impact for the business |')
        md_lines.append('| Core Focus Areas | Execution, stakeholder alignment, metrics |')
        md_lines.append('| Key Competencies | Problem solving, delivery, collaboration |')
        md_lines.append('| Consulting & Client Engagement | Discovery, advisory, influence |')
        md_lines.append('| AI / Analytics / Domain | Practical awareness and usage |')
        md_lines.append('| Cultural Fit | Ownership, clarity, bias for action |')
        md_lines.append('\n---\n')

        # Per-candidate sections
        for r in results:
            name = r['filename']
            md_lines.append('# 🧩 Candidate Summary')
            md_lines.append(f'**Name:** {name}')
            md_lines.append('**Current Role:** —')
            md_lines.append('**Experience:** —')
            md_lines.append('**Industry / Domain Expertise:** —')
            md_lines.append('**Education:** —')
            md_lines.append('**Location:** —')
            md_lines.append(f"**Key Themes / Keywords:** {', '.join(r.get('top_strengths', [])[:6]) or '—'}")
            md_lines.append('\n---\n')

            # Ensure non-empty strengths/gaps
            if not r.get('top_strengths'):
                base = os.path.splitext(name)[0]
                heur = [w.title() for w in base.replace('_',' ').replace('-',' ').split() if len(w) > 2][:3]
                r['top_strengths'] = heur or ['General delivery', 'Stakeholder collaboration']
            if not r.get('key_gaps'):
                r['key_gaps'] = ['No critical gap surfaced']

            md_lines.append('# 📊 Comparative Fit Analysis (JD vs Resume)')
            md_lines.append('| **Dimension** | **Evaluation** | **Commentary** |')
            md_lines.append('|----------------|----------------|----------------|')
            mark = eval_mark(r['match_percentage'])
            sig = ", ".join(r.get("top_strengths", [])[:3]) or '—'
            gap_one = (r.get('key_gaps') or ['—'])[0]
            md_lines.append(f'| Domain Expertise | {mark} | Signals: {sig} |')
            md_lines.append(f'| Consulting & Advisory Orientation | {mark} | Based on profile narrative |')
            md_lines.append(f'| AI / Analytics Awareness | {mark} | Tooling/awareness inferred |')
            md_lines.append(f'| Account Growth / Leadership | {mark} | Team/initiative ownership |')
            md_lines.append(f'| Client Gravitas (C-suite Influence) | {mark} | Stakeholder influence indicators |')
            md_lines.append(f'| Communication & Storytelling | {mark} | Clarity of outcomes |')
            md_lines.append(f'| Technical or Delivery Depth | {mark} | Depth vs breadth balance |')
            md_lines.append(f'| Cultural Fit (Consulting + Innovation) | {mark} | Bias for action, collaboration |')
            md_lines.append('')

            md_lines.append('# 💪 Key Strengths')
            strengths = r.get('top_strengths', [])[:5] or ['General delivery', 'Collaboration']
            for s in strengths:
                md_lines.append(f'- {s}')
            md_lines.append('')

            md_lines.append('# ⚠️ Gaps / Risks')
            md_lines.append('| Gap | Explanation | Impact |')
            md_lines.append('|------|-------------|---------|')
            gaps = r.get('key_gaps', [])[:3] or ['No critical gap surfaced']
            for g in gaps:
                md_lines.append(f'| {g} | — | Medium |')
            md_lines.append('')

            md_lines.append('# 🧾 Scorecard Summary')
            def to_star(score):
                # Map 0-100 to 1-5
                return max(1, min(5, round(score/20)))
            star = to_star(r['match_percentage'])
            md_lines.append('| Category | Rating (1–5) | Comment |')
            md_lines.append('|-----------|--------------|----------|')
            for cat in ['Domain Fit','Consulting Gravitas','AI / Analytics Awareness','Account Growth Leadership','Client Relationship / Communication','Cultural Fit']:
                md_lines.append(f'| {cat} | {star} | Derived from resume signals |')
            overall10 = round(r['match_percentage']/10, 1)
            verdict = '✅ Strong Fit' if r['match_percentage']>=75 else ('⚠️ Partial Fit' if r['match_percentage']>=55 else '❌ Not a Fit')
            md_lines.append('')
            md_lines.append(f'**Overall Fit Score:** {overall10} / 10  ')
            md_lines.append(f'**Verdict:** {verdict}')
            md_lines.append('\n---\n')

            md_lines.append('# ✅ Final Recruiter Verdict')
            md_lines.append('> Candidate shows relevant capability signals with room to validate consulting gravitas and delivery depth. Recommend next-step screening focused on stakeholder influence, structured problem solving, and measurable impact.')
            md_lines.append('')

        # Summary comparison table
        md_lines.append('# ⚖️ Multi-Candidate Comparison Table')
        header = '| **Criteria** | ' + ' | '.join([r['filename'] for r in results]) + ' |'
        sep = '|---------------|' + '|'.join(['----------------' for _ in results]) + '|'
        md_lines.append(header)
        md_lines.append(sep)
        def row_line(label):
            vals = []
            for r in results:
                vals.append(f"{round(r['match_percentage']/10,1)} / 10")
            return f"| {label} | " + " | ".join(vals) + " |"
        for crit in ['Domain Relevance','Consulting Orientation','AI / Analytics Exposure','Client Growth Aptitude','Cultural Fit','**Overall Fit Score**']:
            md_lines.append(row_line(crit))

        report_markdown = '\n'.join(md_lines)

        return jsonify({'success': True, 'results': results, 'report_markdown': report_markdown})
    except Exception as e:
        logging.error(f"Error in evaluate_batch: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == "__main__":
    # Initialize logging
    logging.basicConfig(level=logging.INFO,
                       format='%(asctime)s - %(levelname)s - %(message)s')

    # Initialize database
    init_db()
    
    # Update database schema
    update_db_schema()
    
    try:
        # Initialize Pinecone safely
        vectorstore = initialize_pinecone()
        
        # Build BM25 index
        logging.info("🔍 Building BM25 index...")
        build_bm25_index(POLICIES_FOLDER)
        
        # Set up LLM and QA chain
        logging.info("🤖 Setting up LLM and QA chain...")
        llm, qa_chain, retriever = setup_llm_chain()
        
        # Start Flask server with ASGI support using hypercorn
        logging.info("🌐 Starting server...")
        from hypercorn.config import Config
        from hypercorn.asyncio import serve

        config = Config()
        config.bind = ["localhost:5000"]
        config.use_reloader = True
        
        asyncio.run(serve(asgi_app, config))
        
    except Exception as e:
        logging.error(f"❌ Startup error: {str(e)}")
        raise
